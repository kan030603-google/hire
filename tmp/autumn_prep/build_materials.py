from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"C:\Users\kan\OneDrive\Desktop\研一\秋招")
OUT = ROOT / "output" / "秋招准备材料"
OUT.mkdir(parents=True, exist_ok=True)

INK = "1F2937"
BLUE = "1F4E79"
BLUE_2 = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F3F4F6"
MID_GRAY = "667085"
LIGHT_GRAY = "D0D5DD"
GREEN = "18794E"
GOLD = "8A6116"
RED = "9B1C1C"
WHITE = "FFFFFF"

A4_WIDTH = 11906
A4_HEIGHT = 16838
GUIDE_MARGIN = 936  # 0.65 in; named A4/CJK override of compact_reference_guide
GUIDE_TABLE_WIDTH = A4_WIDTH - GUIDE_MARGIN * 2 - 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_east_asia_font(element, name: str) -> None:
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def style_font(style, name: str, size: float, color: str = INK, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold
    set_east_asia_font(style.element, name)


def paragraph_border(paragraph, side: str = "bottom", color: str = BLUE, size: int = 10, space: int = 3) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    tag = OxmlElement(f"w:{side}")
    tag.set(qn("w:val"), "single")
    tag.set(qn("w:sz"), str(size))
    tag.set(qn("w:space"), str(space))
    tag.set(qn("w:color"), color)
    pbdr.append(tag)


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_keep(paragraph, keep_next: bool = False, keep_lines: bool = True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        ppr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        ppr.append(OxmlElement("w:keepLines"))


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def configure_section(section, margin_dxa: int = GUIDE_MARGIN) -> None:
    section.page_width = Twips(A4_WIDTH)
    section.page_height = Twips(A4_HEIGHT)
    section.top_margin = Twips(margin_dxa)
    section.bottom_margin = Twips(margin_dxa)
    section.left_margin = Twips(margin_dxa)
    section.right_margin = Twips(margin_dxa)
    section.header_distance = Inches(0.34)
    section.footer_distance = Inches(0.34)


def add_running_furniture(section, short_title: str) -> None:
    hp = section.header.paragraphs[0]
    hp.text = short_title
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.runs[0]
    hr.font.name = "Microsoft YaHei"
    hr.font.size = Pt(8.5)
    hr.font.color.rgb = rgb(MID_GRAY)
    set_east_asia_font(hr._element, "Microsoft YaHei")

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("秋招面试准备  |  ")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(MID_GRAY)
    set_east_asia_font(r._element, "Microsoft YaHei")
    add_field(fp, "PAGE")
    r = fp.add_run(" / ")
    r.font.size = Pt(8)
    add_field(fp, "NUMPAGES")

    # Populate all header/footer variants explicitly. Some locally installed
    # Word templates export alternating pages with even-page parts despite the
    # document-level odd/even flag being off.
    for target in (section.even_page_header, section.first_page_header):
        target.is_linked_to_previous = False
        target_element = target._element
        for child in list(target_element):
            target_element.remove(child)
        for child in section.header._element:
            target_element.append(deepcopy(child))
    for target in (section.even_page_footer, section.first_page_footer):
        target.is_linked_to_previous = False
        target_element = target._element
        for child in list(target_element):
            target_element.remove(child)
        for child in section.footer._element:
            target_element.append(deepcopy(child))


def add_numbering_definition(doc: Document, num_id: int, abstract_id: int, decimal: bool) -> None:
    numbering = doc.part.numbering_part.element
    abs_num = OxmlElement("w:abstractNum")
    abs_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abs_num.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if decimal else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if decimal else "•")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Microsoft YaHei")
    fonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rpr.append(fonts)
    lvl.append(rpr)
    abs_num.append(lvl)
    # OOXML requires all abstractNum definitions to precede concrete num
    # instances. Appending an abstractNum after an existing num makes Word
    # repair the numbering part and can silently turn bullets into a running
    # decimal list.
    first_num_index = next(
        (i for i, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abs_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)


def apply_num(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)


def setup_guide_doc(short_title: str) -> Document:
    doc = Document()
    configure_section(doc.sections[0])
    add_running_furniture(doc.sections[0], short_title)
    styles = doc.styles
    style_font(styles["Normal"], "Microsoft YaHei", 10.5)
    normal = styles["Normal"].paragraph_format
    normal.space_before = Pt(0)
    normal.space_after = Pt(6)
    normal.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 25, INK, 0, 5),
        ("Subtitle", 11.5, MID_GRAY, 0, 13),
        ("Heading 1", 16, BLUE_2, 18, 10),
        ("Heading 2", 13, BLUE_2, 14, 7),
        ("Heading 3", 11.5, BLUE, 10, 5),
    ]:
        style_font(styles[name], "Microsoft YaHei", size, color, bold=name != "Subtitle")
        pf = styles[name].paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.keep_with_next = True
        pf.keep_together = True

    if "Code Block" not in [s.name for s in styles]:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    style_font(code, "Microsoft YaHei", 9, INK)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.12

    add_numbering_definition(doc, 41, 41, decimal=False)
    add_numbering_definition(doc, 42, 42, decimal=True)
    return doc


def add_title_block(doc: Document, kicker: str, title: str, subtitle: str, meta: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(kicker.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(BLUE_2)
    set_east_asia_font(r._element, "Microsoft YaHei")
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_border(p, "bottom", BLUE, 12, 5)
    p = doc.add_paragraph(subtitle, style="Subtitle")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(meta)
    r.font.size = Pt(9.5)
    r.font.color.rgb = rgb(MID_GRAY)
    set_east_asia_font(r._element, "Microsoft YaHei")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    if level == 1:
        paragraph_border(p, "bottom", LIGHT_GRAY, 7, 3)


def add_para(doc: Document, text: str, bold_prefix: str | None = None, color: str | None = None,
             italic: bool = False, after: float | None = None) -> None:
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        set_east_asia_font(r._element, "Microsoft YaHei")
        r = p.add_run(text[len(bold_prefix):])
    else:
        r = p.add_run(text)
    if color:
        for run in p.runs:
            run.font.color.rgb = rgb(color)
    if italic:
        for run in p.runs:
            run.italic = True


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, 41)
        p.add_run(item)


def add_steps(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, 42)
        p.add_run(item)


def add_callout(doc: Document, label: str, text: str, kind: str = "info") -> None:
    fill, line, label_color = {
        "info": (PALE_BLUE, BLUE_2, BLUE),
        "note": (PALE_GRAY, MID_GRAY, INK),
        "good": ("EAF7F0", GREEN, GREEN),
        "warn": ("FFF5E6", GOLD, GOLD),
        "risk": ("FDECEC", RED, RED),
    }[kind]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.06)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, fill)
    paragraph_border(p, "left", line, 18, 5)
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.color.rgb = rgb(label_color)
    set_east_asia_font(r._element, "Microsoft YaHei")
    p.add_run(text)
    set_keep(p)


def add_flow(doc: Document, lines: Sequence[str]) -> None:
    p = doc.add_paragraph(style="Code Block")
    shade_paragraph(p, PALE_GRAY)
    paragraph_border(p, "left", BLUE_2, 14, 5)
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(8.8)
        set_east_asia_font(r._element, "Microsoft YaHei")


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int], indent: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[idx]))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int] | None = None) -> None:
    cols = len(headers)
    if widths is None:
        base = GUIDE_TABLE_WIDTH // cols
        widths = [base] * cols
        widths[-1] += GUIDE_TABLE_WIDTH - sum(widths)
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = value
        set_cell_shading(cell, PALE_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.2)
                set_east_asia_font(r._element, "Microsoft YaHei")
    trpr = table.rows[0]._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(value)) > 14 else WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)
                    set_east_asia_font(r._element, "Microsoft YaHei")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_qa(doc: Document, qa_items: Sequence[tuple[str, str, Sequence[str] | None]]) -> None:
    for idx, (question, answer, followups) in enumerate(qa_items, start=1):
        add_heading(doc, f"Q{idx}. {question}", 2)
        add_para(doc, "建议回答：" + answer, bold_prefix="建议回答：")
        if followups:
            add_bullets(doc, ["追问展开：" + x for x in followups])


def add_sources(doc: Document, sources: Sequence[str], boundary: str) -> None:
    add_heading(doc, "资料依据与表达边界", 1)
    add_callout(doc, "边界", boundary, "warn")
    add_para(doc, "主要依据：")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run("；".join(sources))
    r.font.size = Pt(8.3)
    set_east_asia_font(r._element, "Microsoft YaHei")


def save_doc(doc: Document, filename: str) -> Path:
    path = OUT / filename
    doc.core_properties.author = "阚海"
    doc.core_properties.title = path.stem
    doc.core_properties.subject = "2027 届秋招准备材料"
    doc.core_properties.keywords = "秋招, 大模型应用, Agent, 面试"
    doc.save(path)
    return path


# ---------------------------- Resume ----------------------------

def setup_resume_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    configure_section(section, margin_dxa=792)  # 0.55 in; old one-page resume-inspired override
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.25)
    styles = doc.styles
    style_font(styles["Normal"], "Microsoft YaHei", 9.1)
    styles["Normal"].paragraph_format.space_after = Pt(2.2)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    for name, size, color, before, after in [
        ("Heading 1", 13.5, INK, 7, 4),
        ("Heading 2", 10.8, BLUE, 4, 2),
        ("Heading 3", 9.8, INK, 2, 1),
    ]:
        style_font(styles[name], "Microsoft YaHei", size, color, True)
        styles[name].paragraph_format.space_before = Pt(before)
        styles[name].paragraph_format.space_after = Pt(after)
        styles[name].paragraph_format.keep_with_next = True
    add_numbering_definition(doc, 51, 51, decimal=False)

    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("内容扩展版｜投递前请删除所有【待确认】字段并按岗位压缩")
    r.font.size = Pt(7.5)
    r.font.color.rgb = rgb(MID_GRAY)
    set_east_asia_font(r._element, "Microsoft YaHei")
    return doc


def add_resume_header(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("阚 海")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = rgb(INK)
    set_east_asia_font(r._element, "Microsoft YaHei")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("大模型应用 / Agent 算法 / 搜索推荐方向")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = rgb(BLUE)
    set_east_asia_font(r._element, "Microsoft YaHei")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("13046859853  |  2025282050209@whu.edu.cn")
    r.font.size = Pt(9.3)
    set_east_asia_font(r._element, "Microsoft YaHei")


def resume_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph(title, style="Heading 1")
    paragraph_border(p, "bottom", BLUE, 9, 2)


def add_right_tab(paragraph, right_text: str) -> None:
    usable = A4_WIDTH - 792 * 2
    paragraph.paragraph_format.tab_stops.add_tab_stop(Twips(usable), WD_TAB_ALIGNMENT.RIGHT)
    paragraph.add_run("\t" + right_text)


def resume_row(doc: Document, left: str, right: str, center: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.5)
    r = p.add_run(left)
    r.bold = True
    if center:
        r = p.add_run("    " + center)
        r.bold = False
    add_right_tab(p, right)


def resume_project(doc: Document, name: str, role: str, bullets: Sequence[str]) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_after = Pt(1.5)
    p.add_run(name).bold = True
    r = p.add_run(f"  |  {role}")
    r.bold = False
    r.font.color.rgb = rgb(MID_GRAY)
    for item in bullets:
        p = doc.add_paragraph()
        apply_num(p, 51)
        p.paragraph_format.space_after = Pt(1.6)
        p.paragraph_format.line_spacing = 1.04
        p.add_run(item)


def build_resume() -> Path:
    doc = setup_resume_doc()
    add_resume_header(doc)
    resume_section(doc, "教育背景")
    resume_row(doc, "武汉大学｜测绘工程｜硕士", "2025.09 - 2027.06")
    add_bullets_resume(doc, [
        "研究方向：基于 YOLO 的车载影像敏感信息识别；基于 GAN 的地理数据多尺度更新",
        "课程：人工智能与机器学习、深度学习与大模型",
    ])
    resume_row(doc, "中南大学｜测绘工程｜学士", "2021.09 - 2025.06")
    add_bullets_resume(doc, [
        "研究方向：基于三维高斯泼溅的无人机影像三维重建",
        "课程：Python、C++、数据结构、科学计算与 MATLAB",
    ])

    resume_section(doc, "实习经历")
    resume_row(doc, "阿里巴巴｜【部门待确认】", "【起止时间待确认】", "大模型应用 / 算法实习生【岗位待确认】")
    resume_project(doc, "EduBot 多 Agent 路由与任务编排", "核心开发 / 链路审计", [
        "参与教育大模型多 Agent 控制链建设，将 Scene 直达、Main ToolCall Handoff、active sticky 与 SubAgent Handback 收敛到统一 Router；明确 Gate advice 与 Router effect 边界，支撑 Main / 专业 Agent 跨轮连续承接。",
        "设计“高精度规则优先 + 单 token 轻量模型 + 双阈值 margin”连续性判定，结合请求级实验配置快照与 fail-open 降级，在时延、误切换风险和任务连续性之间取平衡。",
        "建设 active / paused Task、Topic、Owner、LGI 与 revision/CAS 状态链，完善 PREEMPT / RESUME、ToolCall-ToolResult 闭合及 Chat/Live/SSE 交付语义；梳理并发重叠、迟到结果与游标回退风险。",
        "完成线上 on 桶 163 个多轮 Session / 902 请求全链路对账与 18 个 Session 语义审计，验证终态 898/902 正常并定位恢复规则双向失准、LGI 非零回退与 conflict 无重评等问题；单日实验初步观察到 Session 转讲万物率 +14.9pp、单轮退出率 -4.7pp、放弃率 -2.5pp（1128/197 Sessions，真实干预率 6.7%，需结合小样本口径解释）。",
    ])
    resume_project(doc, "教育对话主动服务", "算法与工程设计", [
        "设计“上一轮异步 prepare、下一轮实时 confirm”的两阶段主动服务链：dialogue_turn Event 触发 L0-L4 Graph 生成 pending steering，下一轮 Delivery Gate 基于最新 Query 再判 fit，仅在 Prompt 开关开启且实际执行 Main 时注入。",
        "构建 L0 状态重建、L1 规则准入/频控、L2 用户姿态感知、L3 Goal+Guidance 联合决策与确定性 post-correction、L4 暂存分发，覆盖危机/拒绝阻断、冷启动、通道 cooldown、结构校验与 quiet turn。",
        "面向多 Worker / 多机房设计连接复用、Redis 水合、TTL、短锁与异步多写；识别非原子消费、candidate goal 跨 Worker 丢失及 200ms/400ms deadline 错配，给出 atomic claim、delivery token、lease 与 ACK 的演进方案。",
        "A/B 结果：On/Off 分别 16,238/15,889 Sessions，实际注入 417 Sessions（2.57%）；平均 Main 轮次 +0.480、多轮率 +1.91pp、深度对话率 +1.66pp、平均会话时长 +12.67 秒，均达到统计显著（p≤0.0096）。",
    ])

    resume_row(doc, "高途教育集团｜用户研发部｜算法研发组", "2026.02 - 2026.04", "大模型应用开发实习生")
    resume_project(doc, "智能搜索对话系统", "核心成员", [
        "参与 C 端主 App 搜索算法链路重构，覆盖 Query 标准化、实体证据收集、意图分类、动态检索路由、混合召回/重排和会话管理；将纯关键词、向量与 LLM 能力按成本和确定性分层。",
        "建设 BERT 意图识别训练链路：融合约 1.4 万真实样本、实体标注与模板合成数据，使新增老师/课程等实体通过词典更新即可泛化；以类别权重、Focal Loss、Label Smoothing 改善长尾类和 hard case，并采用高置信本地模型、低置信轻量 LLM 回退。",
        "重构上下文管理：近期轮次用 role-aware messages 保真，FactMemory 全量注入，EventSummary 通过 RAG 按需召回；结合 tiktoken、轮次上限与超预算摘要控制上下文成本。",
    ])
    resume_project(doc, "用户画像与长期记忆", "负责人", [
        "从 0 到 1 设计离线画像/长期记忆 Pipeline，将多源行为统一为 HistoryEvent，经 session/天/周窗口生成 EventSummary，再增量抽取 FactMemory，并构建 L1 静态、L2 统计、L3 动态兴趣画像。",
        "设计 FactMemory ADD/SKIP/MERGE/REPLACE 生命周期协议：规则前置减少不必要 LLM 调用，按 category+tag 合并并保留 inactive 演化历史、来源与时效；基于 SHA-256 Prompt 缓存、分类重试和 Token 分步追踪控制成本。",
        "处理 40 万+搜索记录、322 万+课程行为、124 万+订单、562 万+社区互动等数据；课程内容关联覆盖率 98.7%，32 个固定记忆标签覆盖率 90%+，EventSummary 线程池并行实现约 4 倍加速。",
    ])
    resume_project(doc, "多模态标签服务", "负责人", [
        "负责教育图文/视频内容的兴趣与地域标签主链，统一图片压缩、视频均匀抽 20 帧、封面优先、16k 音频转写与多模态消息编排；地域标签按“目标受众所在地”而非普通地名抽取建模。",
        "用候选标签、定义和正负样例约束多模态 LLM 输出，结合 JSON 解析、白名单映射、必含/排除词、优先级覆盖、去重和 Top3 限制，降低自创标签与过度泛化。",
        "抽象 mock / OpenAI-compatible Provider，采用 ASR 懒加载、串行锁、单媒体失败不中断和本地文件支持，将算法主链从 HTTP/Apollo/回调基础设施中解耦为可运行 Demo。",
    ])

    resume_section(doc, "相关技能")
    add_bullets_resume(doc, [
        "语言与工程：Python、SQL、C++（了解）、Linux、Git、Docker；Redis、HTTP/SSE、并发与分布式一致性基础",
        "大模型应用：Agent / Handoff、Context Engineering、Long-term Memory、RAG、Prompt Engineering、Tool Calling、结构化输出",
        "算法与框架：PyTorch、Transformer、BERT 意图分类、混合检索/BM25/向量检索、多模态 LLM、ASR；了解 SFT/PEFT 与强化学习",
        "其他：CET-6；具备基础搜索推荐架构与 A/B 实验分析能力",
    ])

    resume_section(doc, "论文及知识产权")
    add_bullets_resume(doc, [
        "ISPRS 2026 Congress (TC II), Abstract Accepted：Image-Assisted Aerial LiDAR Completion with Morphology-Guided Gaussian Ellipsoids（Co-author；重建/数据 Pipeline、实验与可视化）",
        "《地理空间信息》投稿（2025.11，Under Review）：《土地利用数据综合的改进型 Pix2Pix 深度学习模型》（Co-author；数据清洗、改进点梳理）",
        "软件著作权（受理）：《基于深度学习的土地利用数据多尺度更新软件（MSLC-GAN）》（Co-author；数据清洗、软件封装）",
    ])
    add_callout_resume(doc, "投递前待确认：阿里部门、岗位名称与起止时间；论文/软著最新状态；内部数据与项目名是否允许对外披露。")
    return save_doc(doc, "00_阚海_秋招简历_内容扩展版.docx")


def add_bullets_resume(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, 51)
        p.paragraph_format.space_after = Pt(1.2)
        p.paragraph_format.line_spacing = 1.03
        p.add_run(item)


def add_callout_resume(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    shade_paragraph(p, "FFF5E6")
    paragraph_border(p, "left", GOLD, 14, 3)
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = rgb(GOLD)
    set_east_asia_font(r._element, "Microsoft YaHei")


# ----------------------- Project 1: Alibaba Router -----------------------

def build_ali_router() -> Path:
    doc = setup_guide_doc("阿里｜EduBot 多 Agent 路由与任务编排")
    add_title_block(
        doc,
        "PROJECT INTERVIEW BRIEF 01",
        "EduBot 多 Agent 路由与任务编排",
        "从“每轮选 Agent”升级为可持续、可暂停恢复、可审计的多轮任务控制面",
        "公司：阿里巴巴  |  角色：核心开发 / 链路审计【具体岗位待确认】  |  状态：一期已上线 + 目标架构设计",
    )
    add_callout(doc, "30 秒版本", "我做的是教育大模型的多 Agent 控制链。核心不是把 Query 分成 Main 或 SubAgent，而是让专业任务在多轮里有稳定身份、能被抢断和恢复，并在并发、流式输出和异常情况下保持状态可解释。我参与收敛 Gate、Router、Task/Topic、ToolCall/Handback 链路，并用生产日志完成 163 个多轮 Session、902 请求的全链路审计，定位恢复判定和 LGI 单调性问题。", "good")

    add_heading(doc, "一、项目背景与业务问题", 1)
    add_para(doc, "Main Agent 适合开放问答、澄清和能力协调，SubAgent 适合持续讲解、辅导或练习。如果每轮都让 Main 重新选择，时延与成本上升且专业任务容易摇摆；如果只用 active_agent 粘住，又无法表达换题、暂停、恢复、完成和同 Agent 多任务。因此需要一个独立控制面管理“本轮归谁”和“跨轮任务如何变化”。")
    add_table(doc, ["方案", "优势", "核心缺陷"], [
        ["每轮 Main 重新选择", "灵活", "重复模型成本；相邻轮抖动；进度难恢复"],
        ["只记 active_agent", "低时延", "无法区分任务身份；易错误黏住；不能暂停/恢复"],
        ["Router + Task 生命周期", "连续、可恢复、可审计", "需要状态机、并发控制和协议闭合"],
    ], [2100, 2800, GUIDE_TABLE_WIDTH - 4900])

    add_heading(doc, "二、我的工作与可交付结果", 1)
    add_bullets(doc, [
        "参与统一请求链：把 Scene Direct、Main ToolCall Handoff、active sticky、SubAgent Handback 统一收敛到 Router 循环。",
        "参与 Handoff Gate 的规则优先与单 token 小模型方案，明确 advice、apply 与 final target 三层事实。",
        "围绕 Session / Turn / Topic / Task / Owner / LGI / revision 建立状态语义，覆盖 CREATE、CONTINUE、PREEMPT、RESUME、COMPLETE。",
        "梳理 ToolCall-ToolResult 闭合、Chat/Live/SSE 输出、异常 healing、CAS 与迟到结果等工程边界。",
        "基于生产 SLS 做 on 桶全量链路对账、分支定向抽样与人工语义复核，定位 P0/P1 问题并给出可落地方案。",
    ])
    add_callout(doc, "不要说过头", "一期已经具备 active/paused 两个互斥槽位、关键转换 CAS 和统一 Router；Task Registry、多 paused、真实 checkpoint/restore、command ledger 与 transactional outbox 是目标设计或算法 Demo 合同，不能表述为已完整上线。", "risk")

    add_heading(doc, "三、业务链路", 1)
    add_flow(doc, [
        "用户 Query",
        "  -> 入口解析 / 安全审核 / 确定性干预",
        "  -> PromptGet + Handoff Gate + Proactive Gate 并行准备",
        "  -> Router 校验实验开关、revision 与最新 Task 状态",
        "  -> MAIN | SCENE_DIRECT | ACTIVE_HANDOFF | HANDOFF_INTENT",
        "  -> Main / SubAgent 流式执行",
        "  -> ToolCall / Final / Handback / Error 再进入 Router",
        "  -> SSE 交付 + 消息/状态持久化 + 观测闭环",
    ])
    add_heading(doc, "五种任务动作", 2)
    add_table(doc, ["动作", "触发", "状态变化", "LGI 规则"], [
        ["CREATE", "Scene 首轮或 Main 新建 Handoff", "Main -> 新 active；清旧 paused", "新任务从初始值开始"],
        ["CONTINUE", "当前 Query 仍属 active Task", "保持 Task/Topic/Owner", "吸收合法 ACK"],
        ["PREEMPT", "换题或明确暂离", "active -> paused；本轮 Main", "先保存最新确认值"],
        ["RESUME", "明确恢复且目标唯一", "paused -> active", "恢复 paused 保存值"],
        ["COMPLETE", "SubAgent Handback", "清任务并交回 Main", "完成后不可恢复"],
    ], [1250, 2500, 3000, GUIDE_TABLE_WIDTH - 6750])

    add_heading(doc, "四、核心算法设计", 1)
    add_heading(doc, "4.1 规则优先 + active continuity 小模型", 2)
    add_bullets(doc, [
        "高精度规则先处理明确继续、形式调整、系统控制、暂停/退出、情绪支持等后果清晰的表达。",
        "有 active 且规则未命中时，小模型只输出 0/1，不生成业务答案；输入仅保留当前 Query、上一轮用户话术、任务目标、Topic、最近进度和 Agent 能力。",
        "用 margin = log P(Main) - log P(SubAgent) 做双阈值：明显继续、明显回 Main、中间不确定区。配置示例为 <= -0.5 继续，>= 2.875 抢断，中间 CLARIFY。",
        "异常时如果已确认合法 active，优先 CONTINUE_ACTIVE；状态本身读失败则不做写操作。这里的 fail-open 是用户连续性开放，观测必须标记 fallback。",
    ])
    add_heading(doc, "4.2 paused-only 恢复与为什么后来发现它有问题", 2)
    add_para(doc, "一期只有一个 paused 候选，恢复规则尝试综合 Task ID、目标/摘要关键词、Agent 别名和 resume cue。设计初衷是降低误恢复，但线上审计表明纯词面规则既会在“继续聊/接着告别”等语境中误召，又会漏掉真正的短句“继续”和同主题内容追问。更合理的方向是规则召回 + 语义复核，并用 Query 与 paused goal/topic 的匹配分数选择目标。")

    add_heading(doc, "五、工程设计", 1)
    add_heading(doc, "5.1 权责与状态一致性", 2)
    add_bullets(doc, [
        "Gate 只产 advice；Router 校验实验配置、revision 和目标 Task 后才提交 effect；Task Store 不理解语义。",
        "请求级配置做快照，避免 Gate、Router、Main 在同一请求里读取不同版本。",
        "关键 PREEMPT/RESUME 使用 revision + CAS，复合状态在 Redis 事务中写入；dispatch 前再次确认 active Task。",
        "所有终态应绑定 session_id + task_instance_id + command_id + expected_version，避免同 Agent 的旧结果污染新 Task。",
    ])
    add_heading(doc, "5.2 协议闭合与流式交付", 2)
    add_bullets(doc, [
        "Main 交给 SubAgent 是 ToolCall；任务完成或抢断回 Main 时必须补同 tool_call_id 的 ToolResult，且新 Query 位于闭合结果之后。",
        "active 期间 ToolCall 暂时 orphan 是合法开放协议；Handback、Preempt、过期 healing 时必须闭合。",
        "Task LGI 是任务续传位置，Session 高水位保证全会话帧单调；两者都不是 Agent checkpoint。",
        "Chat SSE、Live SSE 与 JSON 共用 Router，但 final/error/handoff 的外部表达不同，需要分别验收。",
    ])
    add_heading(doc, "5.3 CAS 还不够", 2)
    add_para(doc, "CAS 防止旧快照覆盖新状态，但不能防重复命令、状态已提交而 dispatch 未发生、SSE 已发而消息未落盘、同 Session 双流和跨 Worker 主动候选重复消费。目标方案是 command ledger + request owner/epoch + checkpoint/restore + transactional outbox + 下游幂等。")

    add_heading(doc, "六、线上审计与实验结果", 1)
    add_table(doc, ["证据", "结果", "面试解释"], [
        ["链路全量核验", "163 多轮 Session / 902 请求；898 正常终态", "证明控制链可对账，不等于语义全对"],
        ["决策", "580 次成功：模型 329 + 规则 251", "规则路径也是 Gate 成功，不能漏算"],
        ["并发", "47/163 Session 有重叠；相邻对重叠 11.4%", "并发是常态；CAS conflict 需重评"],
        ["恢复质量", "规则触发精确 0/4；显式“继续”召回 0/4", "样本小但机制缺陷明确；需语义复核"],
        ["LGI", "5 个 Session、6 次非零 -> 0", "场景入口重发导致；服务端需高水位保护"],
        ["单日 A/B", "转讲万物 +14.9pp；退出 -4.7pp；放弃 -2.5pp", "1128/197 Sessions，真实干预率 6.7%，只作初步证据"],
    ], [1900, 3000, GUIDE_TABLE_WIDTH - 4900])
    add_callout(doc, "指标口径", "实验收益必须同时说出窗口、样本量和真实干预率。线上审计的 18 个语义样本是分支定向抽样，不能把 PASS/PARTIAL/FAIL 分布外推为总体比例。", "warn")

    add_heading(doc, "七、关键问题、复盘与下一步", 1)
    add_table(doc, ["问题", "根因", "改进"], [
        ["恢复规则假阳性/假阴性", "词面 cue 缺少当前活动语境与目标匹配", "规则召回后模型复核；goal/topic 与 Query 相似度；候选消歧"],
        ["LGI 回退", "incoming=0 直接覆盖非零 checkpoint", "max(existing,incoming) 或 CAS 单调；重入单独建模"],
        ["CAS conflict 无重评", "正确建议和错误建议都可能被随机丢弃", "按新 revision 重跑低成本规则；记录冲突轮终态"],
        ["任务碎片化", "CLARIFY 与 Main 工具反复新建/恢复", "统一 CLARIFY 语义；同 Topic 恢复与新建边界"],
        ["状态/dispatch 分裂", "跨系统无事务", "next state + command result + event + outbox 同事务"],
    ], [2200, 3500, GUIDE_TABLE_WIDTH - 5700])

    add_heading(doc, "八、面试高频问答", 1)
    add_qa(doc, [
        ("为什么这不是普通意图分类？", "意图分类只回答当前文本像什么；本项目还要结合已有任务、状态版本和恢复资格决定本轮 owner，并维护跨轮 Task 生命周期、协议与流式进度。", ["用‘换题—临时问答—恢复—完成’四轮例子解释状态变化。"]),
        ("Gate 和 Router 为什么分开？", "Gate 是可替换的算法建议层，Router 是控制权威。分开后可以 shadow、分桶、校准模型而不直接改状态，也能由 Router 做开关、revision、CAS 和最终兜底。", ["Gate hit 不能当业务成功，至少还要看 apply、transition、final target 和用户可见结果。"]),
        ("为什么规则在模型前？", "明确暂停/继续等高精度控制意图不需要模型；规则更快、可解释，在历史或模型异常时仍可工作。模型只处理语义模糊区。", ["规则不能无限扩张，恢复规则的线上失准正说明复杂语义需要模型参与。"]),
        ("为什么用 logprob margin 而不是最大概率？", "margin 直接表达 Main 与 SubAgent 两个动作的相对证据，并支持不对称双阈值：自动继续和自动抢断的风险不同，中间区交 Main 澄清。", ["阈值必须按场景与误判成本校准，不是固定常数。"]),
        ("CLARIFY 为什么难？", "它既可以是瞬时让 Main 问一句，也可以被实现成 durable pause。两种语义会影响任务是否可恢复和是否产生碎片，必须先冻结产品合同再编码状态机。", None),
        ("为什么 active 和 paused 不能同时存在？", "这是一期容量简化，不是理论要求。它降低状态组合和恢复消歧成本，但新任务会覆盖旧 paused 的恢复资格。审计显示多数回归指向最近任务，容量暂时可用，短板主要在判定。", None),
        ("Topic 和 Task 有什么区别？", "Topic 是语义上下文边界，Task 是某个 Agent 对该目标的一次执行。一个 Topic 最终可以关联多次 Task；一期 topic_id==task_id 只是迁移简化。", None),
        ("LGI、progress summary、checkpoint 有什么区别？", "LGI 回答客户端播放到哪里，summary 回答用户看到了什么，checkpoint 回答 Agent 内部执行到哪个可恢复安全点。只有 checkpoint 能真正恢复内部工作流。", None),
        ("为什么 CAS 不够？", "CAS 只保护基于旧版本的写，不解决相同 command 重放、状态与 dispatch 原子性、迟到结果、双流和 SSE/持久化分叉，所以还需要幂等台账、owner/epoch、outbox 和 settled。", None),
        ("ToolCall 为什么必须闭合？", "Main 的模型历史要求每个 ToolCall 有配对 ToolResult，否则会认为工具仍未完成，导致重复调用或上下文非法。闭合还必须与最新 Query 保持正确顺序。", None),
        ("线上审计怎么做？", "先按 matched.edu_main_gate 精确取 on 请求，再以 session_id+request_id 关联 Gate、apply、dispatch、transition、终态六类标记做全量对账；异常分支定向抽样，合成整段时间线做语义复核。", ["说明分支抽样不能外推总体比例。"]),
        ("你发现的最关键问题是什么？", "暂停恢复规则双向失准和 LGI 非零回退。前者导致任务误拉回或长期悬置，后者破坏续播单调性；两者都通过全量链路和语义样本相互验证。", None),
        ("如果重做恢复，你会怎么设计？", "先召回合法 paused 候选，再用 Query、goal、topic、recent progress 做语义匹配；规则只处理显式 Task ID 和高精度 cue；低置信度交 Main 澄清，并记录候选、评分、apply 与实际恢复效果。", None),
        ("怎么降低首 token 时延？", "并行 PromptGet/Gate/Proactive，规则前置，active sticky 绕过 Main，Gate 单 token+短 deadline+连接池；但只取消最终不需要的分支，不能把投机 hint 当已提交 route。", None),
        ("目标架构为什么需要 outbox？", "状态提交和跨服务 dispatch 无法放进同一个远程事务。把 dispatch record 与 next state 同事务写入，再由 worker 至少一次投递、下游按 command_id 幂等，覆盖进程中途崩溃窗口。", None),
    ])

    add_heading(doc, "九、两分钟项目陈述模板", 1)
    add_para(doc, "项目面向教育大模型的多轮专业任务承接。早期方案要么每轮让 Main 重新选择 Agent，成本高且不连续；要么只记 active_agent，换题后容易粘住，也无法安全暂停恢复。我们把它抽象成 Gate + Router + Task 状态链：Gate 用高精度规则和单 token 小模型给出继续、抢断、恢复或澄清建议；Router 结合请求级实验配置和 revision 做最终决策，并维护 active/paused Task、Topic、Owner 和 LGI。Main 通过 ToolCall 发起 Handoff，SubAgent 通过 Handback 结束任务，Router 补 ToolResult 形成闭环。工程上关键是把 advice 和 effect 分开、用 CAS 防旧快照覆盖，并分别处理 Chat/Live/SSE 交付。我还做了生产日志审计：对 163 个多轮 Session、902 请求完成全链路对账，终态 898 个正常，同时发现暂停恢复规则精确率和召回都存在机制性问题、5 个 Session 出现 LGI 回退。对应方案是让恢复引入目标语义匹配、对 LGI 做高水位保护，并在 CAS conflict 后按新版本低成本重评。")

    add_sources(doc, [
        "阿里实习材料/01-路由与Handoff完整业务逻辑.md",
        "阿里实习材料/02-任务编排与Topic完整设计.md",
        "阿里实习材料/04-统一架构与链路工程设计.md",
        "阿里实习材料/edubot-on-bucket-audit-2026-08-26.md",
        "阿里实习材料/各类实验收益.txt",
    ], "“当前一期”“算法 Demo”“目标形态”须分开陈述；线上审计结论可以说，未落地的 Task Registry、checkpoint/outbox 只能说设计或演进方案。")
    return save_doc(doc, "01_项目面试_阿里_EduBot多Agent路由与任务编排.docx")


# ----------------------- Project 2: Alibaba Proactive -----------------------

def build_ali_proactive() -> Path:
    doc = setup_guide_doc("阿里｜教育对话主动服务")
    add_title_block(
        doc,
        "PROJECT INTERVIEW BRIEF 02",
        "教育对话主动服务",
        "把高成本理解前移到后台，用下一轮最新 Query 做实时确认，再安全注入 Main",
        "公司：阿里巴巴  |  角色：算法与工程设计【具体岗位待确认】  |  状态：steering 主链已具备，可靠交付/学习闭环待完善",
    )
    add_callout(doc, "30 秒版本", "我设计的是教育对话里的主动引导，不是弹窗推荐。上一轮回答结束后异步用画像和会话跑 L0-L4，生成一个 pending steering；下一轮用户发来最新问题时，再用轻量 Delivery Gate 判断旧候选是否仍合适。只有配置开启且 Router 实际启动 Main 才注入。这样把两次主模型放到后台，实时路径只承担一次轻量 fit 判断，同时用规则、频控和 post-correction 保证主动内容不打断显式需求。", "good")

    add_heading(doc, "一、业务目标与边界", 1)
    add_bullets(doc, [
        "目标：在不打断主任务的前提下，基于画像、历史、掌握度和当前话题发现高价值延伸机会。",
        "不是最终 Router：不能决定 Main/SubAgent owner，也不创建 Handoff Task。",
        "不是最终回答：steering 是 Main 可忽略的内部建议，最新用户 Query 永远优先。",
        "不是每轮必推：正常 miss、quiet turn 和异常时无主动内容都是正确结果。",
        "不是已完成的在线学习系统：当前缺 durable delivery ledger、完整归因、蒸馏与回滚闭环。",
    ])

    add_heading(doc, "二、我的工作", 1)
    add_bullets(doc, [
        "设计 prepare-confirm 两阶段架构，拆分后台 Graph 与下一轮 Delivery Gate，解决候选跨轮过时问题。",
        "设计 L0-L4 分层：状态重建、规则准入、用户姿态感知、Goal+Guidance 联合决策、候选暂存/分发。",
        "把模型输出转成可执行候选：危机/拒绝阻断、cooldown、通道可用性、知识点去重、结构与内容红线校验。",
        "完成 EduBot 侧集成边界设计：Prompt 配置、Handoff Gate、Proactive Gate 并行；Main 注入四条件；失败不阻塞主回答。",
        "梳理多 Worker / 多机房下的 Redis 水合、短锁与异步多写，定位重复消费、goal 丢失和 deadline 状态分叉。",
        "结合 A/B 结果评估真实收益，并明确 Gate fit、Prompt 注入、语义消费和最终效果之间不可跳步。",
    ])

    add_heading(doc, "三、端到端业务链路", 1)
    add_flow(doc, [
        "第 n 轮结束",
        "  -> EduBot 异步 POST /events（门铃，不携带完整真值）",
        "  -> 主动服务重读画像/Session -> L0 -> L1 -> L2 -> L3 -> L4",
        "  -> PendingSteering 写入进程内单槽 + Redis 镜像",
        "第 n+1 轮 Query 到来",
        "  -> 同步 POST /gate（约 200ms 调用侧预算）",
        "  -> 水合 + detach -> 轻量 0/1 fit judge -> fit 返回 steering/behavior",
        "  -> Router 决定本轮 target",
        "  -> 仅 Prompt 开关 on + steering 非空 + 实际执行 Main 时注入",
        "  -> Main 独立回答，可融合/忽略/ToolCall Handoff",
    ])
    add_callout(doc, "核心思想", "候选生成时的上下文属于上一轮；下一轮用户可能换题、拒绝或进入 SubAgent。主动服务必须让最新 Query 否决旧候选，因此不能在生成后立即展示。", "info")

    add_heading(doc, "四、L0-L4 算法设计", 1)
    add_table(doc, ["层", "模型成本", "职责", "失败/阻断语义"], [
        ["L0 ContextBuilder", "0", "重读画像与历史，重建描述状态、turn、handoff", "部分数据可降级；不从空 event 编造"],
        ["L1 RuleGate", "0", "kill switch、cold start、cap、双通道 cooldown", "全部通道不可用则 quiet，零模型成本"],
        ["L2 Observer", "1 次主模型", "描述 topic、act、engagement、opportunity、情绪/危机", "解析失败静默结束；危机/拒绝确定性阻断"],
        ["L3 Reasoner", "1 次主模型", "联合决策 engage、goal action、guidance、channel", "post-correction 删除非法 intent；无 surviving intent 不推进 goal"],
        ["L4 Dispatcher", "0", "投影 Main view，生成 pending，写单槽/Redis", "quiet 不写空候选，不 tick 频控"],
        ["Delivery Gate", "1 次轻量判断", "用最新 Query 判旧候选 fit/unfit", "error/timeout 对主动内容保守关闭"],
    ], [1450, 1250, 4200, GUIDE_TABLE_WIDTH - 6900])

    add_heading(doc, "4.1 为什么 L2 只描述、不决策", 2)
    add_para(doc, "把 Perception 和决策分离可以复用感知结果，并让危机、拒绝、stop/shift 等红线由确定性代码复查。L2 只说明用户状态和机会强弱；L3 才结合 goal、频控、历史干预和可用通道决定是否介入。")
    add_heading(doc, "4.2 为什么 L3 联合决策 Goal 与 Guidance", 2)
    add_para(doc, "如果分别让多个模型定目标、写引导、选通道，容易产生方向冲突和额外成本。一次联合决策能保持 should_engage、goal_action、guidance 与 card_intent 一致，再由代码按状态机和安全规则校正。")
    add_heading(doc, "4.3 post-correction", 2)
    add_bullets(doc, [
        "合法化 create/maintain/advance/revise/branch/drop；不存在 goal 时不能凭空 advance。",
        "should_engage=false 或 drop 时删除 guidance；cooldown 通道 intent 强制删除。",
        "删除未开启 card phase 的幻觉输出；知识点冲突、方向不一致、必填字段缺失时降级 quiet。",
        "拦截危机、健康焦虑、越界工具主题；没有 surviving intent 时不允许推进 goal。",
    ])

    add_heading(doc, "五、工程设计", 1)
    add_heading(doc, "5.1 关键路径与 fail-open", 2)
    add_bullets(doc, [
        "两次主模型在上一轮后台运行，下一轮只做轻量 fit 判断；Event HTTP 不延长上一轮回答。",
        "客户端按 Worker 复用连接，入口并发准备 PromptGet、Handoff Gate 与 Proactive Gate。",
        "对主回答 fail-open：主动服务超时/Redis/模型异常时仍正常回答；对主动内容 fail-closed：未经确认不展示。",
        "Main 消费合同要求核心回答独立完整、最多自然融合一个短延伸、不得暴露内部策略。",
    ])
    add_heading(doc, "5.2 多 Worker / 多机房缓存", 2)
    add_bullets(doc, [
        "进程内 SessionContext 是同 Worker 快路径；Redis 镜像解决 Event 在 A、Gate 在 B 的水合。",
        "Event 后台向多个机房并发写，Gate 默认只读本地，实时路径避免跨区 RTT；Redis 故障退化为 miss。",
        "Gate 只在短锁内 detach 槽位，模型 judge 放在锁外，避免实时请求等待后台模型。",
        "当前 GET-judge-DELETE 不是全局原子，两个 Worker 可重复判断同一候选；多机房异步删除也有短暂重复窗口。",
    ])
    add_heading(doc, "5.3 提交点和 deadline 错配", 2)
    add_para(doc, "当前服务在 Gate fit 时就提交 candidate goal、tick clock 并删除候选，但这早于 EduBot 收到、Prompt 开关、Main 注入和用户可见。更危险的是调用侧约 200ms、服务 Judge 约 400ms：调用方可能已超时继续主回答，服务端却稍后把候选当作已消费。目标是 STAGED->CLAIMED->JUDGED->DELIVERED_ACKED/RELEASED 的原子状态机，并传播 deadline 或感知取消。")

    add_heading(doc, "六、实验收益与正确解读", 1)
    add_table(doc, ["指标", "On", "Off", "提升", "p-value"], [
        ["Session 数", "16,238", "15,889", "+349", "-"],
        ["实际注入 Session", "417（2.57%）", "0", "+417", "-"],
        ["平均 Main 轮次", "3.310", "2.830", "+0.480", "0.000008"],
        ["多轮率（>=2）", "27.26%", "25.36%", "+1.91pp", "0.000105"],
        ["深度对话率（>=5）", "11.66%", "10.01%", "+1.66pp", "0.000002"],
        ["平均会话时长", "176.54s", "163.87s", "+12.67s", "0.0096"],
    ], [2300, 1800, 1800, 1900, GUIDE_TABLE_WIDTH - 7800])
    add_callout(doc, "怎么解释", "这是 on/off 的意向治疗口径，只有 2.57% Session 实际注入。不能把整体 lift 直接解释成“每次注入的因果效果”，还需检查随机化、曝光归因、守护指标和异质性。面试时先给结果，再主动说这一限制，可信度会更高。", "warn")

    add_heading(doc, "七、现状、技术债与演进", 1)
    add_table(doc, ["当前已有", "关键缺口", "目标方案"], [
        ["Event + full fetch", "缺 turn/version/high-watermark 因果锚点", "durable inbox；数据版本写入 decision snapshot"],
        ["内存 + Redis pending", "跨 Worker 非原子 claim；多机房重复窗口", "Lua/CAS atomic claim + delivery token + lease"],
        ["Gate fit 延迟提交", "仍早于真实注入；goal 字段水合不完整", "candidate goal 完整持久化；Injection/User-visible ACK"],
        ["Feedback API / Reflector 局部能力", "缺 durable ledger、归因与策略回滚", "intervention ledger -> 离线评分/审计 -> 版本化蒸馏 -> A/B/回滚"],
        ["Card 状态机框架", "主端 receiver、展示/点击反馈未闭环", "幂等 receiver + 展示 ACK + attribution"],
    ], [3000, 3300, GUIDE_TABLE_WIDTH - 6300])

    add_heading(doc, "八、面试高频问答", 1)
    add_qa(doc, [
        ("为什么要两阶段，而不是上一轮直接推？", "候选基于上一轮上下文，下一轮用户可能换题、拒绝或进入 SubAgent。后台 prepare 降低实时成本，实时 confirm 让最新 Query 拥有否决权。", None),
        ("为什么 Event 只是门铃？", "完整画像和会话保存在权威数据源，Event 只携带身份和触发信息可降低耦合、重复时也能收敛到同一视图；代价是要处理数据落盘可见性和版本锚点。", None),
        ("L1 为什么不做情绪关键词规则？", "危机、拒绝和情绪需要语境，简单关键词误判高。L1 只做确定性的开关/频控准入，L2 模型感知后再用代码红线复查。", None),
        ("为什么 L2 失败要 quiet？", "主动介入属于附加能力，错误干预的风险高于没有干预。解析失败时伪造默认 Perception 会为了产出率牺牲安全。", None),
        ("为什么 Goal 不能在 L4 生成时就提交？", "候选还可能过期或被下一轮否决。若提前推进 goal，系统会误以为已经教过。至少要延迟到交付定义成立，理想是注入或用户可见 ACK。", None),
        ("Gate fit 为什么还不等于成功？", "后面还有网络返回、Prompt 开关、Router 是否执行 Main、Prompt 是否注入、Main 是否采用、最终 responder 和用户效果。每层只能证明自己的阶段。", None),
        ("为什么 steering 0 就丢弃，而 card 可以保留？", "steering 生成相对便宜且强依赖最近语境，旧建议反复尝试更危险；card 是较昂贵成品，fit=0 可能只是时机不对，TTL 内可继续等待。", None),
        ("Redis 为什么需要，为什么又不够？", "Redis 让跨 Worker 看见候选，但普通 GET/DELETE 不提供全局 claim；两个 Worker 仍可同时消费，多机房副本还会短暂重复，所以要有原子状态机和 delivery token。", None),
        ("如何设计 atomic claim？", "用版本比较或 Lua 把 STAGED 原子转 CLAIMED，写 request_id、delivery_token 和 lease deadline；judge 后转 FIT/UNFIT，收到下游 ACK 才 DELIVERED，否则释放或 lease 超时回收。", None),
        ("调用侧 200ms、服务侧 400ms 有什么问题？", "调用方先超时后服务端仍可能 detach/提交候选，造成用户没收到但服务认为已消费。要传播 deadline、支持取消，或让 caller budget 覆盖网络/排队/judge/响应余量。", None),
        ("怎样证明 Main 真用了 steering？", "把 intervention_id 贯穿 Gate 响应、Prompt 组装、Main 执行与 final responder；对 Main 文本/ToolCall 做语义消费判定，并和最新 Query、最终可见内容、后续行为关联。", None),
        ("A/B 结果怎么讲才严谨？", "先给样本量、实际注入率和四个显著 lift，再说明它是 on/off 总体口径，不等于 treated-only 因果效应；需要随机化检查、守护指标、曝光日志和分人群分析。", None),
        ("频控为什么按通道分开？", "steering 和 card 的成本、打扰程度与交付机制不同，应该各有 count 和 last_output_turn；候选生成/miss 不消耗额度，真实交付才 tick。", None),
        ("怎么保证最新用户意图优先？", "Delivery Gate 用最新 Query 重判，Main 的 behavior contract 要求先完整回答显式需求；危机、拒绝、shift 信号阻断，且 SubAgent 路由不注入 steering。", None),
        ("如果让你继续做，P0 是什么？", "先修交付状态一致性：完整持久化 candidate goal，原子 claim + lease + ACK，统一 deadline；然后补 intervention ledger，才能可靠频控、归因和策略学习。", None),
    ])

    add_heading(doc, "九、两分钟项目陈述模板", 1)
    add_para(doc, "这个项目做的是教育对话中的主动引导，但我们不把它当即时推荐。因为上一轮结束时生成的建议，到下一轮可能已经过时，所以我把链路设计成 prepare-confirm：上一轮结束后 EduBot 异步发 Event，主动服务重读画像和 Session，经过 L0 状态重建、L1 规则准入、L2 用户姿态感知、L3 Goal 与 Guidance 联合决策、L4 暂存，得到 pending steering；下一轮用户发来最新 Query 时，再由轻量 Delivery Gate 做 0/1 fit 判断。只有请求配置开启、Gate 返回非空且 Router 实际启动 Main 时才注入，Main 仍要先完整回答用户，必要时可以忽略或转交 SubAgent。工程上我们把两次主模型放到后台，实时路径只做一次轻量判断，并用连接复用、Redis 跨 Worker 水合、TTL 和短锁控制时延。但我也识别到当前 GET-judge-DELETE 不是全局原子、candidate goal 跨 Worker 不完整，以及调用侧 200ms 小于服务侧 400ms 的状态分叉。对应目标是 atomic claim、delivery token、lease 和下游 ACK。实验中 On/Off 为 16,238/15,889 Sessions，实际注入率 2.57%，平均 Main 轮次、多轮率、深度对话率和时长都有统计显著提升；我会把它解释为总体 on/off 结果，而不是把所有 lift 都直接归因到被注入样本。")

    add_sources(doc, [
        "阿里实习材料/03-主动服务完整业务与算法设计.md",
        "阿里实习材料/04-统一架构与链路工程设计.md",
        "阿里实习材料/各类实验收益.txt",
        "阿里实习材料/ec1cfe27bcc0e4264ce0bf7c53ee1d6d.jpg",
    ], "可以说 steering 生成与实时 Gate 主链已具备；完整可靠交付、Card 接入、durable feedback/reflector 和在线学习闭环仍是缺口。实验图展示总体 on/off 指标，treated-only 因果归因尚不足。")
    return save_doc(doc, "02_项目面试_阿里_主动服务系统.docx")


# ----------------------- Project 3: Gaotu Search -----------------------

def build_gaotu_search() -> Path:
    doc = setup_guide_doc("高途｜智能搜索对话系统")
    add_title_block(
        doc,
        "PROJECT INTERVIEW BRIEF 03",
        "智能搜索对话系统",
        "Query 标准化、实体证据、分层意图识别、动态检索路由与多层上下文",
        "公司：高途教育集团  |  部门：用户研发部算法研发组  |  角色：核心成员  |  时间：2026.02 - 2026.04",
    )
    add_callout(doc, "30 秒版本", "我参与了高途 C 端主 App 搜索链路重构。系统先做输入标准化和实体证据收集，再用高置信 BERT、低置信轻量 LLM 的分层意图识别决定是检索还是生成；检索侧按意图和实体构造动态 Filter，并用当前 Query、上下文和画像做乘法重排。我的重点还包括 BERT 训练链路和多层上下文：近期轮次保真、长期 FactMemory 注入、历史 EventSummary 按需 RAG，解决多轮指代与 token 成本。", "good")

    add_heading(doc, "一、项目背景", 1)
    add_table(doc, ["用户输入", "难点", "系统处理"], [
        ["白马老师有课吗", "别名映射", "词典证据 -> 老师实体 -> 课程查询"],
        ["baima laoshi", "拼音输入", "实体词典生成拼音变体并最长匹配"],
        ["考研数学怎么学", "无显式实体的知识问题", "低置信回退 LLM / 生成链"],
        ["他有没有网课", "跨轮指代", "context entity + role-aware messages"],
        ["AI 闪学在哪", "产品功能模糊表达", "证据 + 产品功能意图 -> 单结果入口"],
    ], [2700, 3000, GUIDE_TABLE_WIDTH - 5700])
    add_para(doc, "纯关键词无法覆盖别名、拼音和语义变体；纯向量召回容易混淆找老师/找课程；纯 LLM 每请求调用时延和成本过高。因此采用本地快速处理优先、证据驱动和按不确定性升级。")

    add_heading(doc, "二、我的工作", 1)
    add_bullets(doc, [
        "参与六步算法主链重构：QueryProcessor -> EvidenceCollector -> Intent -> ResponseBuilder -> Retriever/LLM -> SessionManager。",
        "负责 BERT 意图识别训练链路：真实样本、实体增强、模板×实体、同义词与多轮上下文增强；WeightedTrainer 支持类别权重、Focal Loss、Label Smoothing。",
        "设计 BERT 高置信优先、轻量 LLM 低置信回退、规则仅演示兜底的分层架构，按文本/图片设置差异阈值。",
        "负责上下文管理：把历史从普通字符串改为 role-aware messages，精确 token 计数，近期保真、超限摘要，并与用户画像/长期记忆联动。",
        "参与基于意图和实体的动态 VectorDB Filter、当前/上下文/画像权重重排与检索/生成路由。",
    ])
    add_callout(doc, "生产与 Demo 边界", "技术文档来自对原项目算法主链的去基础设施化 Demo。生产服务还包含 Django、Apollo、Eureka、SkyWalking、数据库治理和上下游编排；Demo 可独立运行并自动降级，但不能把 Demo 本身说成完整生产服务。", "warn")

    add_heading(doc, "三、整体链路", 1)
    add_flow(doc, [
        "Raw Query",
        "  -> 全角转半角 -> 繁简转换 -> 实体词典驱动的拼音转汉字 -> 空白清理 -> 分词",
        "  -> Evidence：精确匹配 -> token 匹配 -> 模糊/向量回退 -> 去重融合",
        "  -> Intent：BERT(高置信) -> 轻量 LLM(低置信) -> 演示规则兜底",
        "  -> ResponseBuilder：实体 target/filter 萃取 + 只基于当前 Query 的意图覆盖",
        "  -> 检索类：动态 Filter + 混合召回 + 乘法重排；生成类：LLM 回答",
        "  -> SessionManager：保存本轮并更新 context entity / 多层上下文",
    ])

    add_heading(doc, "四、算法设计", 1)
    add_heading(doc, "4.1 证据收集", 2)
    add_bullets(doc, [
        "精确匹配使用多模式词典；文档采用 FlashText/Trie 思路，适合大量实体别名一次扫描。",
        "token 匹配降低置信度，模糊匹配只作回退；证据按 (type,value) 去重保留最高分。",
        "实体证据在意图识别前产生，因为 BERT 本身不知道‘白马老师=褚佳麟’，外置知识能帮助分类。",
        "推理时用与训练端一致的实体类型标记，最长匹配避免长短别名相互覆盖。",
    ])
    add_heading(doc, "4.2 BERT 分层意图识别", 2)
    add_table(doc, ["层", "条件", "作用"], [
        ["本地 BERT", "纯文本 confidence >=0.85；带图 >=0.90", "10-50ms 量级，先处理确定样本"],
        ["轻量 LLM", "BERT 低置信或不可用", "64-token JSON 意图；处理长尾/歧义"],
        ["规则兜底", "Demo 无 BERT/LLM", "保持主链可运行；生产不依赖这一层"],
    ], [1700, 3600, GUIDE_TABLE_WIDTH - 5300])
    add_para(doc, "训练数据约包括 1.4 万真实用户标注、1.4 万实体打标版、5.4 万模板×实体合成和约 1.9 万均衡合并版。核心不是死记实体名，而是学习“实体类型 + 句式 -> 意图”，新增实体只需更新词典即可泛化。")
    add_heading(doc, "4.3 类别不均衡与 hard case", 2)
    add_bullets(doc, [
        "类别权重解决频率不均衡，让长尾类在梯度中不被多数类淹没。",
        "Focal Loss 解决难度不均衡，降低 easy sample 权重，强调边界样本；gamma=2 是默认配置。",
        "Label Smoothing=0.1 缓解歧义样本过拟合并改善置信度校准，使阈值回退更可信。",
        "以 weighted F1 选 checkpoint，连续 3 个 epoch 不提升 early stop；同时应关注各类 recall/F1，而不只看 accuracy。",
    ])
    add_heading(doc, "4.4 动态检索与重排", 2)
    add_para(doc, "只对课程、老师、产品功能等检索类意图访问向量库。Filter 先加 intent_word 和 status=1，再把当前实体、上下文实体和画像作为可选召回约束。多个实体用 OR 保召回，再在排序阶段按当前实体×3、上下文×1.5、画像×1.2 做乘法增益；乘法能避免低基础相关性的垃圾结果被固定加分抬高。")

    add_heading(doc, "五、上下文工程", 1)
    add_table(doc, ["层", "内容", "注入方式", "原因"], [
        ["FactMemory", "稳定、浓缩的长期事实", "system context 全量", "量少、信息密"],
        ["EventSummary", "历史 Session 摘要", "按当前 Query RAG", "数量随历史增长，多数不相关"],
        ["当前 Session 近期轮次", "原始 user/assistant 消息", "messages 保真", "支撑指代与连续对话"],
        ["当前 Session 早期溢出", "超过预算的旧轮次", "在线摘要", "防 context overflow"],
    ], [1900, 2800, 2200, GUIDE_TABLE_WIDTH - 6900])
    add_bullets(doc, [
        "原实现把全部历史拼成普通字符串，角色信息丢失、token 估算粗糙、成本随轮次线性增长。",
        "改为从最新轮向前累加，默认 2000-token 预算、保留近 5 轮，超预算部分用与离线 EventSummary 对齐的 Prompt 压缩。",
        "Session 另有最多 10 轮与 token 上限的双重防线；tiktoken 不可用时才用字符估算。",
        "V1 用户画像是离线 per-user JSON，尚非在线 API；EventSummary embedding/RAG 是计划接入条件。",
    ])

    add_heading(doc, "六、关键工程取舍", 1)
    add_table(doc, ["取舍", "为什么", "边界"], [
        ["BERT 优先、LLM 回退", "确定样本低延迟低成本，歧义样本保质量", "需要校准置信度和覆盖率"],
        ["实体知识外置", "新实体更新词典即可生效", "词典覆盖不足时退回句式特征/LLM"],
        ["OR Filter + 乘法重排", "先保召回再保精度", "权重需离线/在线实验，不能长期手调"],
        ["可选依赖降级", "便于 Demo 移植和故障隔离", "能力退化必须显式记录"],
        ["近期保真 + 历史摘要", "兼顾指代精度与 token 成本", "摘要可能损失细节，需要 RAG 补充"],
    ], [2500, 4000, GUIDE_TABLE_WIDTH - 6500])

    add_heading(doc, "七、面试高频问答", 1)
    add_qa(doc, [
        ("为什么不用纯 LLM 做意图？", "大量高频 Query 可以被本地模型稳定处理，纯 LLM 增加 7-9s 级 API 延迟、成本和可用性风险；高置信 BERT + 低置信 LLM 能把复杂样本留给更强模型。", ["7-9s 和 10-50ms 是文档中的量级/估算，不要说成个人线上压测结论。"]),
        ("新增老师为什么可以不重训？", "训练和推理都把实体替换/标记为统一类型 token，模型学的是‘老师实体 + 句式’，实体表更新后 EvidenceCollector 自动识别并加标记，权重无需记住具体名字。", None),
        ("实体标注会不会依赖词典？", "会，所以混合保留约 30% untagged 样本，让未命中词典时仍可依靠句式；低置信再回退 LLM。", None),
        ("类别权重、Focal Loss、Label Smoothing 分别解决什么？", "类别权重解决类频率不均，Focal Loss 解决样本难度不均，Label Smoothing 解决硬标签和过度自信；三者互补但要用消融验证。", None),
        ("为什么 weighted F1 仍不够？", "它仍受大类权重影响，可能掩盖长尾类。线上路由更应看各类 precision/recall、低置信覆盖率、LLM 回退率和业务代价矩阵。", None),
        ("阈值 0.85/0.90 怎么定？", "应在验证集做可靠性图/温度缩放，结合误判成本、LLM 成本与覆盖率选择；图片场景不确定性更高所以阈值更保守。", None),
        ("为什么证据收集放在意图前？", "实体类型本身是强意图特征，能把业务知识从模型权重中外置；先有 evidence 再分类，再用 intent 指导检索。", None),
        ("为什么 Filter 用 OR？", "AND 容易因实体不一致召回为空；OR 先保留老师相关和学科相关候选，再用当前实体/上下文/画像的差异权重重排。", None),
        ("为什么用乘法重排？", "乘法保留基础相关性门槛，低相关结果再多业务 boost 也不容易被抬到顶部；加法可能让垃圾结果靠固定加分越级。", None),
        ("上下文为什么不能全量拼接？", "角色丢失、token 线性增长、旧信息干扰当前意图。分层后近期轮次保真，长期事实全量，历史摘要按需召回，超预算再压缩。", None),
        ("FactMemory 和 EventSummary 有何不同？", "FactMemory 是跨 Session 的稳定事实，量少可全量注入；EventSummary 是每次会话摘要，数量大、按 Query 相关性召回。", None),
        ("如何评估系统？", "离线看实体识别、意图 per-class F1/校准、召回 NDCG/Recall@K；在线看回退率、首包、零结果率、搜索成功/点击/转化、多轮承接和成本，并按意图分桶。", None),
        ("有什么没有做完？", "技术资料没有给出完整线上意图准确率和消融结果；V1 画像是离线 JSON，EventSummary 在线向量召回尚是接入规划；Demo 的规则兜底不代表生产。", None),
    ])

    add_heading(doc, "八、两分钟项目陈述模板", 1)
    add_para(doc, "高途 C 端搜索输入很杂，包括老师别名、拼音、产品功能、无实体知识问答和跨轮指代。纯关键词覆盖不足，纯向量难区分意图，纯 LLM 又有成本和时延。我们把链路拆成 Query 标准化、Evidence、Intent、检索路由和 Session 上下文。先通过全半角、繁简、拼音和分词清洗 Query，再用精确、token、模糊等多路证据识别业务实体；BERT 处理高置信样本，低置信回退轻量 LLM。BERT 训练上，我使用真实用户样本与实体增强，让模型学习实体类型而不是名字，并用类别权重、Focal Loss 和 Label Smoothing 处理长尾与 hard case。检索类意图按意图和实体构造动态 Filter，多个条件先 OR 保召回，再对当前 Query 实体、上下文实体和画像做乘法重排。多轮上下文方面，把原先的历史字符串改为 role-aware messages：FactMemory 作为稳定事实全量注入，EventSummary 按需 RAG，当前 Session 近期轮次保真，超预算部分摘要。这样整套方案把确定性、本地模型和 LLM 按成本分层，同时明确 Demo 与生产基础设施的边界。")

    add_sources(doc, [
        "高途实习材料/阚海_简历_27应届.pdf（职责与项目名称的第一参考）",
        "高途实习材料/TECH_aisearch.md（算法与工程细节的第二参考）",
        "高途实习材料/TECH_userprofile.md（长期记忆与上下文联动）",
    ], "旧简历是职责表述的第一依据；技术文档中的性能数字部分是估算/推荐配置，未给出完整线上效果，不应包装成个人实测收益。独立目录是原生产项目的算法 Demo。")
    return save_doc(doc, "03_项目面试_高途_智能搜索对话系统.docx")


# ----------------------- Project 4: Gaotu Profile -----------------------

def build_gaotu_profile() -> Path:
    doc = setup_guide_doc("高途｜用户画像与长期记忆")
    add_title_block(
        doc,
        "PROJECT INTERVIEW BRIEF 04",
        "用户画像与长期记忆",
        "从多源行为到可演化 FactMemory，再到静态、统计与动态兴趣画像",
        "公司：高途教育集团  |  角色：负责人  |  时间：2026.02 - 2026.04  |  形态：V1 离线构建 + 本地展示",
    )
    add_callout(doc, "30 秒版本", "我从 0 到 1 做了教育场景用户画像和长期记忆离线系统。先把搜索、课程、订单、社区和服务通话统一成 append-only HistoryEvent，再按 session/天/周聚合成 EventSummary，按周抽取并用 ADD/SKIP/MERGE/REPLACE 协议维护 FactMemory；画像分 L1 静态、L2 统计和 L3 最近兴趣。工程上做了规则前置、Prompt Hash 缓存、错误分类、Token 追踪和 EventSummary 并行，既控制成本又保留可追溯的记忆演化。", "good")

    add_heading(doc, "一、项目背景与目标", 1)
    add_para(doc, "在线教育的用户信息分散在搜索、课程浏览、订单、社区互动、客服通话等系统。直接把原始行为塞给 LLM 会产生数据格式不一致、噪声大、成本高、事实重复/冲突、时效难治理等问题。本项目先建立统一中间层，再把稳定事实和动态兴趣分开建模，为搜索个性化和长期对话提供结构化输入。")
    add_table(doc, ["层", "回答的问题", "更新特点"], [
        ["HistoryEvent", "用户发生了什么最小行为？", "append-only；保留事件时间与来源"],
        ["EventSummary", "一个 session/一天/一周发生了什么？", "规则+LLM 窗口聚合；去噪压缩"],
        ["FactMemory", "哪些事实值得长期记住？", "按 category+tag 增量演化，有 active/inactive"],
        ["UserProfile", "稳定属性、统计强度、近期兴趣是什么？", "L1 规则、L2 统计、L3 LLM"],
    ], [1700, 4300, GUIDE_TABLE_WIDTH - 6000])

    add_heading(doc, "二、我的工作", 1)
    add_bullets(doc, [
        "定义 HistoryEvent / EventSummary / FactMemory / UserProfile Schema 与端到端 Pipeline。",
        "为搜索、课程、订单、社区、服务通话设计过滤、时间解析、窗口聚合和规则/LLM 分工。",
        "设计 FactMemory 固定两级标签、TTL 与 ADD/SKIP/MERGE/REPLACE merge 协议，保留来源和 inactive 历史。",
        "实现 L1 静态画像、L2 统计画像与 L3 最近四周兴趣画像，并设计搜索服务接入方式。",
        "建设缓存、重试、内容风控、Token/耗时 Trace、线程池并行和 Demo 数据集。",
    ])

    add_heading(doc, "三、完整 Pipeline", 1)
    add_flow(doc, [
        "多源 Parquet / 服务通话",
        "  -> [Rule] 清洗、过滤、时间统一 -> HistoryEvent",
        "  -> [Rule + LLM] session / 天 / 周聚合 -> EventSummary",
        "  -> [Rule] L1 Static + L2 Stat（含学科参与强度）",
        "  -> [LLM] 按周抽取候选 FactMemory",
        "  -> [Rule 前置 + LLM] ADD / SKIP / MERGE / REPLACE",
        "  -> [LLM] 最终 FactMemory + 最近 4 周摘要 -> L3 InterestProfile",
        "  -> DemoUserBundle JSON（构建与展示分离）",
    ])
    add_heading(doc, "数据规模与处理边界", 2)
    add_table(doc, ["数据", "规模", "主要用途/过滤"], [
        ["AI 搜索", "400,431 条", "按 session 摘要；过滤无教育内容/风控失败"],
        ["课程行为", "3,222,414 条", "近 91 天；有效 course_number 约 11.6%"],
        ["订单", "1,248,432 笔", "按 ISO 周规则聚合；关联课程维表"],
        ["社区互动", "5,621,869 条", "近 91 天；纯曝光跳过，有效互动约 1.7%"],
        ["课程维表", "1,035,669 门", "补课程名/学科/年级，摘要覆盖率 98.7%"],
        ["评论互动", "177,844,530 条", "当前不进入 Pipeline，避免误报处理规模"],
    ], [2000, 2100, GUIDE_TABLE_WIDTH - 4100])

    add_heading(doc, "四、不同数据为何采用不同聚合", 1)
    add_table(doc, ["来源", "窗口", "方法", "原因"], [
        ["AI 搜索", "session", "LLM 摘要", "天然对话单元；保留用户问题与上下文"],
        ["课程行为", "天", "LLM + 规则 fallback", "浏览频繁，需要按日压缩并关联课程内容"],
        ["订单", "ISO 周", "纯规则", "金额/退款/课程是结构化事实，LLM 无增益"],
        ["社区互动", "天", "LLM + 规则 fallback", "需要结合动态文本判断教育主题"],
        ["服务通话", "一通一条", "规则拼装 22 个结构字段", "ASR 后端已结构化，避免重复 LLM"],
    ], [1600, 1500, 3000, GUIDE_TABLE_WIDTH - 6100])

    add_heading(doc, "五、FactMemory 算法", 1)
    add_heading(doc, "5.1 为什么按 category + tag 分组", 2)
    add_para(doc, "Merge 不应在全量记忆中任意比较，而是先按 category+tag 找 active 同组事实，缩小候选并保持业务含义一致。V1 固定 32 个 tag，基于 5 个种子用户 995 条 FactMemory 设计，覆盖率 90%+；长尾可进入兜底治理，但不能让 tag 无限膨胀。")
    add_heading(doc, "5.2 Merge 协议", 2)
    add_table(doc, ["动作", "适用情况", "状态处理"], [
        ["SKIP", "文本相同或已有事实完全覆盖", "不新增；完全相同由规则直接判定"],
        ["ADD", "同组为空或新事实独立", "新增 active；同组为空不调 LLM"],
        ["MERGE", "新旧互补", "同组旧 active 全部 inactive；生成合并事实"],
        ["REPLACE", "新旧冲突/状态变化", "同组旧 active 全部 inactive；新事实继承时效"],
    ], [1500, 4300, GUIDE_TABLE_WIDTH - 5800])
    add_bullets(doc, [
        "规则前置减少成本：完全相同直接 SKIP，同组为空直接 ADD，只有同组非空且文本不同才调 LLM。",
        "MERGE/REPLACE 将全部同组旧记忆设为 inactive，并合并 source_summary_ids，保留演化链。",
        "created_at 取事件窗口时间而非 Pipeline 运行时间，避免离线补跑污染时间语义。",
        "TTL 由候选继承；trait 类型倾向 MERGE，状态变化才谨慎 REPLACE。",
    ])

    add_heading(doc, "六、画像设计", 1)
    add_table(doc, ["画像层", "实现", "示例", "为什么这样做"], [
        ["L1 StaticProfile", "用户宽表规则映射", "年龄/地区/年级/稳定偏好", "确定性强，不需要 LLM"],
        ["L2 StatProfile", "订单/活跃/课程/标签统计", "付费强度、AI 使用、学科参与度", "可计算、可解释"],
        ["L3 InterestProfile", "最终 FactMemory + 最近 4 周摘要", "近期兴趣、学习诉求、成长需求", "动态语义复杂，适合 LLM"],
    ], [1800, 3000, 2700, GUIDE_TABLE_WIDTH - 7500])

    add_heading(doc, "七、工程设计与收益", 1)
    add_bullets(doc, [
        "构建/展示分离：build_all.py 调 LLM 产出 per-user JSON；Streamlit 只读 JSON，零 LLM、秒级展示。",
        "Prompt 缓存键 = SHA-256(model + system + user prompt) 前 16 位，Prompt 或模型变化自动 miss。",
        "最多重试 2 次：网络、408/429/5xx 退避；内容风控和其他客户端错误不重试。",
        "TokenTracker 区分等效总 token、实际 API token、缓存调用，并按 event_summary/fact_extract/fact_merge/interest 分步统计。",
        "AI 搜索与社区 EventSummary 用 ThreadPoolExecutor(max_workers=4) 并行，文档记录约 4 倍加速。",
        "课程内容感知摘要覆盖率 98.7%；32 个固定 FactMemory tag 覆盖率 90%+。",
    ])
    add_callout(doc, "V1 边界", "当前是离线 JSON，无在线 API；只聚焦 5 个种子用户构建，50 个种子子集用于数据准备。服务化、Kafka、FactMemory embedding 与 EventSummary 在线 RAG 是规划，不是已上线。", "risk")

    add_heading(doc, "八、下游如何接 AI 搜索", 1)
    add_table(doc, ["产物", "建议接入", "原因"], [
        ["FactMemory", "system context 全量注入", "量少、长期稳定、信息密度高"],
        ["L1/L2/L3 Profile", "结构化 system context", "直接描述用户稳定属性和近期兴趣"],
        ["EventSummary", "embedding 后按 Query RAG", "数量随历史增长，全量注入成本高且噪声大"],
    ], [2000, 3100, GUIDE_TABLE_WIDTH - 5100])

    add_heading(doc, "九、面试高频问答", 1)
    add_qa(doc, [
        ("为什么要三层记忆？", "HistoryEvent 保原子事实和溯源，EventSummary 做窗口压缩，FactMemory 承载可演化长期事实。直接从原始行为到长期记忆会让成本、噪声和更新逻辑耦合。", None),
        ("为什么订单不用 LLM？", "金额、支付、退款、课程等字段结构化且规则可解释，LLM 增加成本和不确定性；只有需要理解文本语义的搜索/社区/课程内容才使用 LLM。", None),
        ("为什么按周 merge？", "周是 V1 的成本和时效折中：降低每事件调用量，同时保留用户状态变化顺序。更实时可改天级或事件流，但需要增量服务和幂等。", None),
        ("为什么 EventSummary 不做衰减？", "旧摘要在当期已经提炼进 FactMemory，后续只用于相关性 RAG；FactMemory 的 MERGE/REPLACE 承担事实演化，RAG 本身又按相关性筛选。", ["这依赖‘旧摘要不反复参与 fact 抽取’的增量合同。"]),
        ("如何处理冲突事实？", "在同 category+tag 内让 LLM 判 REPLACE 或 MERGE；REPLACE 把旧 active 置 inactive并保留来源，不能直接删除，否则无法审计演化。", None),
        ("为什么规则前置？", "完全相同和无同组是确定性情形，直接 SKIP/ADD 可省大量 LLM 调用，也降低随机性。", None),
        ("TTL 怎么设计？", "明确时间锚点的备考阶段、年级、目标等可设 expires_at，再结合类别 base TTL；泛化表达宁可不设，避免伪精确。更新时由新候选继承/覆盖。", None),
        ("怎么避免 LLM 缓存污染？", "缓存 key 包含模型、system prompt 和 user prompt，内容任何变化都产生新 key；还要隔离环境/租户、记录版本，敏感内容不应长期明文缓存。", None),
        ("线程池为什么能 4 倍？", "EventSummary 的多个 session/天任务主要等待外部 LLM I/O，线程池能并发隐藏等待；如果是 CPU 密集型 Python 计算，GIL 下线程未必线性加速。", None),
        ("如何评估记忆质量？", "抽样评估事实正确性、完整性、冲突处理、时效、可溯源性和下游任务增益；离线可看 tag 覆盖/merge action 分布，在线看检索命中、回答一致性和个性化收益。", None),
        ("如何做隐私治理？", "源数据最小化、身份与内容分级、Prompt/日志脱敏、缓存 TTL、访问控制、删除传播；长期记忆要允许查看、纠正和删除，敏感属性限定用途。", None),
        ("V1 最大局限是什么？", "离线 per-user JSON 不能支撑在线新鲜度与高并发查询；下一步应服务化、做增量触发和存储索引，同时明确 EventSummary RAG 与 FactMemory 查询 SLA。", None),
    ])

    add_heading(doc, "十、两分钟项目陈述模板", 1)
    add_para(doc, "这个项目解决的是教育用户信息散落在搜索、课程、订单、社区和服务通话中，既无法直接用于个性化，也不适合把原始数据全塞给大模型。我从 0 到 1 设计了三层记忆和三层画像。底层先把不同表转成统一的 append-only HistoryEvent；再按数据语义用 session、天或周聚合成 EventSummary，其中订单等结构化数据走规则，搜索和社区文本走 LLM；最后按周抽取 FactMemory，并用 ADD、SKIP、MERGE、REPLACE 协议维护演化。为了控制成本，同组为空直接 ADD、文本相同直接 SKIP，只有真正冲突或补充才调用 LLM；旧事实不删除而是 inactive，并保留 source ids 和事件时间。画像则分 L1 静态规则、L2 统计和 L3 最近四周动态兴趣。工程上我做了 Prompt Hash 缓存、错误分类重试、Token 分步追踪和 EventSummary 线程池并行。资料记录课程内容关联覆盖率 98.7%、固定 32 标签覆盖率 90%+、并行约 4 倍加速。V1 是离线 JSON，不会把服务化和在线 RAG 说成已上线；下游设计是 FactMemory/Profile 直接注入，历史 EventSummary embedding 后按 Query 召回。")

    add_sources(doc, [
        "高途实习材料/阚海_简历_27应届.pdf（职责与项目名称的第一参考）",
        "高途实习材料/TECH_userprofile.md（实现、规模、收益和 V1 边界）",
        "高途实习材料/TECH_aisearch.md（下游上下文接入关系）",
    ], "可以明确说负责人、从 0 到 1 与离线 V1；不要说在线 API、Kafka、embedding/RAG 已上线。177M 评论数据未参与 Pipeline，不得并入处理量。")
    return save_doc(doc, "04_项目面试_高途_用户画像与长期记忆.docx")


# ----------------------- Project 5: Gaotu Tags -----------------------

def build_gaotu_tags() -> Path:
    doc = setup_guide_doc("高途｜多模态标签服务")
    add_title_block(
        doc,
        "PROJECT INTERVIEW BRIEF 05",
        "多模态标签服务",
        "图文/视频统一媒体编排 + 候选约束的 MLLM 决策 + 本地规则纠偏",
        "公司：高途教育集团  |  角色：负责人  |  时间：2026.02 - 2026.04  |  文档形态：生产主链抽取后的独立 Demo",
    )
    add_callout(doc, "30 秒版本", "我负责教育内容发布后的兴趣和地域标签主链。图文侧做图片读取、压缩和多图编排；视频侧做均匀抽 20 帧、封面优先、16k 音频转写，再转成多帧图片+文本给多模态模型。模型不能自由造标签，只能从候选集合里选，并参考定义和正负样例；返回后还要经过白名单映射、必含/排除词、优先级覆盖、去重和 Top3。工程上做了 Provider 抽象、ASR 懒加载与串行锁、单媒体失败不中断。", "good")

    add_heading(doc, "一、业务问题", 1)
    add_para(doc, "内容发布后需要抽取兴趣标签和地域标签，供上游系统消费和下游分发。兴趣标签描述主题，如物理、中考、地方教育通知；地域标签不是普通 NER，而是判断最应该接收内容的用户当前所在地。输入既有图文也有视频，信息可能分布在标题、封面、正文帧和音频中。")
    add_table(doc, ["难点", "为什么难", "对应方案"], [
        ["多模态异构", "图片、视频时序、封面、音频证据强弱不同", "统一媒体编排，视频转多帧+ASR 文本"],
        ["标签体系受控", "模型容易自创或按字面过度泛化", "候选、定义、正负样例 + 返回白名单"],
        ["业务语义特殊", "地域标签不是提到哪个地名", "Prompt 明确目标受众地域，普通提及返回空"],
        ["工程依赖重", "cv2/moviepy/FunASR/LLM 网关", "懒加载、provider 抽象、可降级 Demo"],
    ], [2200, 3900, GUIDE_TABLE_WIDTH - 6100])

    add_heading(doc, "二、我的工作", 1)
    add_bullets(doc, [
        "负责 PostTagAnalyzer 媒体编排：图文/视频分流、本地/远程读取、图片 resize/base64、视频临时文件与清理。",
        "实现视频均匀抽 20 帧、末帧补齐、封面插到 frames[0]、16k WAV 导出与 FunASR 转写。",
        "组织兴趣/地域候选标签，拍平地域树并去重，构造图文/视频差异化 Prompt 和 OpenAI-compatible 多模态 messages。",
        "实现 LLM JSON 解析、tagNumber/tagDefine 映射、必含/排除词、优先级覆盖、去重和最多 3 标签。",
        "抽象 MockHeuristic / OpenAICompatible Provider，做 ASR 懒加载、audio_lock 和单图片失败不中断。",
    ])

    add_heading(doc, "三、端到端链路", 1)
    add_heading(doc, "3.1 图文", 2)
    add_flow(doc, [
        "image_urls + image_paths",
        "  -> 逐张读取 -> cv2 解码 -> 宽度 >800 等比缩放",
        "  -> JPEG + base64 -> 标题/正文 + 候选标签 + 定义/样例",
        "  -> 多模态 LLM -> JSON -> 白名单映射 -> 规则过滤 -> tags/regionTags",
    ])
    add_heading(doc, "3.2 视频", 2)
    add_flow(doc, [
        "video_url/path -> 远程下载或本地直读 -> moviepy duration/fps",
        "  -> 按总帧数均匀抽 20 帧 -> resize/JPEG/base64 -> 封面插入首位",
        "  -> 有音轨且 enable_asr：导出 16k WAV -> FunASR -> transcript",
        "  -> 标题 + transcript（核心）+ frames（辅助）-> MLLM -> 规则后处理",
    ])

    add_heading(doc, "四、算法设计", 1)
    add_heading(doc, "4.1 为什么视频抽 20 帧", 2)
    add_para(doc, "按总帧数均匀采样能对不同长度视频保持固定成本，并补最后一帧提高尾部覆盖。20 不是理论最优，只是信息量、传输/token 成本和实现复杂度的工程折中；缺点是无时序语义、可能错过关键镜头，后续可用镜头切分或关键帧评分优化。")
    add_heading(doc, "4.2 为什么封面优先", 2)
    add_para(doc, "教育短视频的封面通常浓缩标题与主题。把封面置于第一张可让模型先看到高信息密度证据，但也可能受到标题党影响，因此正文帧和 ASR 文本仍需交叉验证。")
    add_heading(doc, "4.3 受约束标签选择", 2)
    add_bullets(doc, [
        "Prompt 给候选标签列表，不允许模型自由生成；标签定义防止只按字面理解。",
        "正负样例明确边界，例如‘细胞分裂过程详解’可选生物，‘数学家的故事’不等于数学课程内容。",
        "地域标签判断目标受众所在地，提到地点但没有定向分发语义时返回空。",
        "解析阶段只保留候选映射中存在的标签，阻断自创和格式污染。",
    ])
    add_heading(doc, "4.4 为什么模型后还要规则", 2)
    add_para(doc, "Prompt 无法完全保证业务边界。后验规则对学科、生态环境、竞赛活动、地方教育通知等重点标签做必含/排除词和优先级压制；之后去重并限制 Top3。它不是完全规则化，而是让模型负责复杂理解、规则负责高风险边界。")

    add_heading(doc, "五、工程设计", 1)
    add_table(doc, ["设计", "收益", "代价/边界"], [
        ["图片压到 800 宽", "降低带宽和视觉 token", "小字可能损失，需按 OCR/内容类型评估"],
        ["单张图失败继续", "部分容错，不因坏图丢整帖", "需要记录有效图片数与降级原因"],
        ["ASR 懒加载", "图文路径不加载重模型", "首次视频请求有冷启动"],
        ["audio_lock", "规避模型线程安全/显存竞争", "视频 ASR 并发被串行化"],
        ["Mock Provider", "无网关也能跑通结构", "不看视觉内容，不能代表真实效果"],
        ["OpenAI-compatible", "便于切模型/网关", "依赖外部服务和 API key"],
    ], [2300, 3200, GUIDE_TABLE_WIDTH - 5500])

    add_heading(doc, "六、已知问题与改进", 1)
    add_table(doc, ["技术债", "影响", "改进"], [
        ["needHidden 地域未过滤", "隐藏候选可能进 Prompt", "候选构建阶段过滤并加单测"],
        ["标签源本身有噪声", "模型选择空间受污染", "标签数据治理、版本与质量校验"],
        ["keywords/word_freq 未进决策", "有计算但无收益", "进入 Prompt 或规则并做消融"],
        ["mock 不看图", "只能证明链路，不能证明效果", "明确标注能力边界或构造可解释视觉 stub"],
        ["stream=True 与 requests 非流式混用", "协议语义不一致", "统一真实/非流式客户端合同"],
        ["Demo 未标准打包", "CI/迁移不规范", "pyproject.toml + wheel + 依赖 extra"],
        ["均匀抽帧无时序", "漏关键镜头", "shot boundary / CLIP 关键帧 / ASR 时间对齐"],
    ], [2600, 3200, GUIDE_TABLE_WIDTH - 5800])

    add_heading(doc, "七、面试高频问答", 1)
    add_qa(doc, [
        ("为什么不是直接把视频传给模型？", "固定抽帧+ASR 能控制成本、兼容图片消息接口，也让证据可观测；直接视频理解依赖模型能力和更高资源。", None),
        ("20 帧怎么定的？", "它是工程折中，不是最优常数。应通过长短视频分层、标签离线集和成本曲线比较 8/12/20/32 帧，观察 F1/一致性与 token/时延。", None),
        ("均匀抽帧有什么问题？", "没有镜头边界和时序意识，可能漏掉短暂关键画面。可用场景切分、CLIP 相似度去冗余、信息熵/字幕变化选帧，并与 ASR 时间戳对齐。", None),
        ("为什么视频 Prompt 更重视音频？", "教育视频核心知识往往在讲解语音中，画面可能只是板书或素材；图文则标题和图片更关键。权重是业务先验，仍要用消融验证。", None),
        ("地域标签为什么不是 NER？", "NER 只识别内容提到的地点，分发需要判断最适合接收者所在地。‘介绍北京历史’未必只投北京，而‘北京市中考报名通知’具有明确地域定向。", None),
        ("为什么候选约束重要？", "可阻断模型自创标签，保证输出能映射到线上 tagNumber/tagDefine，也让评估空间固定、便于版本治理。", None),
        ("规则会不会越来越多？", "会，因此只治理高风险/高频边界，规则要有版本、命中率和误杀率；低风险长尾仍交模型，定期用错误集决定增删。", None),
        ("为什么需要 Provider 抽象？", "把媒体编排和标签决策接口解耦，支持离线 mock、真实网关和后续模型切换；还能分别测试链路正确性和模型质量。", None),
        ("audio_lock 的取舍？", "锁换取模型线程安全和资源稳定，但吞吐受限。生产可用独立 ASR 服务、GPU worker 池、队列+批处理或每进程模型副本。", None),
        ("如何评估？", "构造多标签人工集，兴趣/地域分别看 micro/macro F1、precision@3、空标签准确率、一致性；视频做标题/封面/帧/ASR 消融，并统计时延、token、失败和降级率。", None),
        ("单图失败为什么不中断？", "帖子可能有多张图，坏一张不应丢弃其余有效证据。返回时需记录有效输入数，全部失败再走纯文本或空结果。", None),
        ("Demo 和原服务差别？", "Demo 移除了 HTTP、Apollo、远端标签拉取、回调和部署，只保留真实算法主链；所以可以讲算法设计，但不能拿 Demo 运行等同生产 SLA。", None),
    ])

    add_heading(doc, "八、两分钟项目陈述模板", 1)
    add_para(doc, "这个项目为教育图文和视频内容抽取兴趣与地域标签。难点一是多模态证据分散，难点二是线上标签体系受控，难点三是地域标签并不等于地名识别。我负责媒体编排和标签决策主链：图文逐张读取并压到 800 宽、编码为多模态消息；视频按总帧数均匀抽 20 帧，补末帧，封面放第一位，有音轨时导出 16k WAV 用 FunASR 转写。模型输入里带候选标签、定义和正负样例，只能从候选集合选择；返回后只保留可映射 tagNumber 的结果，再做必含/排除词、优先级覆盖、去重和 Top3。工程上用 Provider 抽象支持 mock 和 OpenAI-compatible，ASR 懒加载避免图文启动重依赖，audio_lock 换取模型线程安全，单张图失败不中断整帖。复盘中明确了隐藏地域未过滤、标签源噪声、关键词没有真正进入决策、均匀抽帧无时序等问题。面试时我会强调当前文档是从原生产服务抽出的算法 Demo，不把它当完整生产系统。")

    add_sources(doc, [
        "高途实习材料/阚海_简历_27应届.pdf（职责与项目名称的第一参考）",
        "高途实习材料/tagdoc.md（算法主链、工程设计与技术债）",
    ], "可说负责生产项目中的算法主链，并将其独立为 Demo；不可说 Demo 自身包含 HTTP/Apollo/回调或证明了真实多模态效果。资料未提供准确率、时延或线上收益，不新增数字。")
    return save_doc(doc, "05_项目面试_高途_多模态标签服务.docx")


# ----------------------- Tech Handbook -----------------------

def build_tech_handbook() -> Path:
    doc = setup_guide_doc("技术栈八股与项目追问手册")
    add_title_block(
        doc,
        "TECHNICAL INTERVIEW HANDBOOK",
        "技术栈八股与项目追问手册",
        "围绕多 Agent、LLM 应用、检索、记忆、多模态与高并发工程，按“核心答案 + 项目映射”组织",
        "适用：大模型应用 / Agent 算法 / 搜索推荐 / Python 后端面试  |  版本：2027 届秋招初稿",
    )
    add_callout(doc, "使用方法", "先掌握 P0 题并能映射到一个真实项目；P1 题用于二面深挖；P2 题用于系统设计。回答顺序统一为：定义 -> 为什么 -> 项目怎么做 -> 边界/改进。", "info")

    add_heading(doc, "一、准备优先级地图", 1)
    add_table(doc, ["优先级", "主题", "必须映射的项目"], [
        ["P0", "Session/Topic/Task、Handoff、状态机、CAS/幂等/outbox", "阿里 EduBot"],
        ["P0", "Context、Long-term Memory、RAG、token 预算", "高途搜索 + 用户画像"],
        ["P0", "BERT、Transformer、Focal Loss、Label Smoothing、阈值", "高途智能搜索"],
        ["P0", "HTTP/SSE/WebSocket/gRPC、连接池、超时重试、网关", "阿里两项目"],
        ["P1", "Redis、多 Worker、多机房、原子 claim、分布式锁", "阿里主动服务"],
        ["P1", "BM25/向量/混合检索/重排", "高途智能搜索"],
        ["P1", "多模态、抽帧、ASR、受约束输出", "高途多模态标签"],
        ["P1", "A/B、p-value、意向治疗、曝光归因", "阿里实验"],
        ["P2", "可观测性、安全隐私、容量与降级系统设计", "全部项目"],
    ], [1200, 4600, GUIDE_TABLE_WIDTH - 5800])

    chapters = [
        ("二、多 Agent 与控制面", [
            ("Session、Turn、Topic、Task 分别是什么？", "Session 是连续交互外层容器；Turn 是一次输入到稳定输出；Topic 是语义边界；Task 是某 Agent 对某目标的一次有生命周期执行。混用会导致上下文、状态和恢复对象错位。", "阿里项目一期 topic_id==task_id 只是简化，目标是 Topic 1:N Task。"),
            ("Handoff、Handback、ToolCall/ToolResult 如何配合？", "Main 用 ToolCall 表达交出任务的意图，Router 创建/恢复 Task 后 dispatch；SubAgent 完成时 Handback，Router 将任务置终态并补配对 ToolResult，再让 Main 收尾。", "active 时 ToolCall 暂时未闭合是合法；抢断/完成/过期时必须闭合。"),
            ("sticky routing 有什么价值和风险？", "价值是连续任务绕过 Main，降低时延和抖动；风险是用户换题后错误黏住。需要 Gate 对最新 Query 与 active Task 做连续性判断。", "EduBot 用规则优先+小模型，在 Router 最终校验后生效。"),
            ("为什么 Gate 不是 Router？", "算法建议与状态权威分离，便于 shadow/A-B/校准、开关和并发校验。Router 才能基于 revision/CAS 修改 durable owner。", "任何指标都要区分 advice、apply、transition 和 final target。"),
            ("checkpoint 和对话摘要有什么不同？", "摘要是给人/模型看的语义进度；checkpoint 是 Agent 可 restore 的内部执行状态引用。摘要不能恢复工具栈、中间结果或安全点。", "阿里一期只有 summary+LGI，无通用 checkpoint。"),
        ]),
        ("三、状态机、并发一致性与可靠消息", [
            ("CAS/乐观锁解决什么？", "读取时带 version，提交时要求 current_version==expected_version，防旧快照覆盖新状态。适合冲突较少、读多写少的状态。", "EduBot PREEMPT/RESUME 使用 revision/CAS；冲突后仍需重评。"),
            ("CAS 为什么不能代替幂等？", "CAS 防并发覆盖，同一个 command 网络重试仍可能在新版本上重复执行。幂等需要 command_id 台账返回第一次结果。", "目标台账保存 command、target、precondition、result 和 resulting_version。"),
            ("transactional outbox 是什么？", "业务状态和待投递事件在同一数据库事务写入，worker 再至少一次投递，下游按 idempotency key 去重，避免状态已变但消息未发。", "适合 Task 激活后 dispatch、Handback 后 complete 等跨服务窗口。"),
            ("exactly-once 能实现吗？", "端到端通常通过 at-least-once + 幂等效果逼近，而不是依赖网络一次送达。要定义 exactly-once 是消息、状态还是业务副作用。", "主动候选当前不是 exactly-once，需要 atomic claim + delivery token。"),
            ("分布式锁与 lease 有什么区别？", "锁强调互斥，lease 带过期时间和持有者身份，适合崩溃恢复。释放要校验 token，避免旧持有者删掉新锁。", "Pending steering claim 应带 request_id/token/deadline。"),
            ("如何处理迟到结果？", "所有终态带 task_id、command_id、expected_version/epoch；只允许合法 owner 在未 settled 时提交。旧结果记录为 stale，不更新新任务。", "同 Agent 不等于同 Task。"),
        ]),
        ("四、高并发与 Python 服务", [
            ("线程、进程、协程怎么选？", "I/O 密集用协程或线程隐藏等待；CPU 密集 Python 受 GIL 影响，优先多进程/原生库/GPU；混合服务常用 async I/O + 独立推理 worker。", "高途 EventSummary 外部 LLM I/O 用 ThreadPoolExecutor；ASR 可独立 worker 池。"),
            ("什么是背压？", "下游处理不过来时，上游必须限速、排队有界、拒绝或降级，避免无限队列拖垮内存和尾延迟。", "Event 合并、Session 单飞、队列长度和超时都是主动服务的背压手段。"),
            ("高并发下为什么关注 p99？", "均值掩盖排队、GC、慢依赖和锁竞争；用户体验和超时通常由尾部决定。应拆连接、排队、模型、SSE、持久化阶段。", "Gate 总耗时不能简单等同模型耗时。"),
            ("连接池解决什么？", "复用 TCP/TLS 连接，降低握手、端口和 CPU 成本；要配置最大连接、keep-alive、获取超时、空闲回收和 DNS。", "阿里客户端按 Worker 复用连接。"),
            ("超时、重试、熔断、隔离如何组合？", "每跳有 deadline；只重试幂等且可恢复错误，指数退避+jitter；持续失败熔断；线程池/连接池隔离慢依赖；总体预算小于上游 deadline。", "主动 Gate 超时返回空，不阻断主回答；高途 LLM 只重试网络/408/429/5xx。"),
            ("什么是惊群和请求合并？", "同一热点 key 失效时大量请求同时回源形成惊群。可用 singleflight、互斥更新、随机 TTL、stale-while-revalidate。", "同 Session 短时间重复 Event 可合并。"),
        ]),
        ("五、HTTP、SSE、WebSocket、gRPC 与网关", [
            ("HTTP keep-alive 和连接池的关系？", "keep-alive 允许单连接复用，连接池管理多条可复用连接及并发借还。HTTP/2 还能在一条连接上多路复用，但仍需流控和最大并发。", "模型/主动服务客户端都应复用连接。"),
            ("SSE 与 WebSocket 怎么选？", "SSE 是服务端到客户端单向、基于 HTTP、自动重连和文本事件，适合 LLM 流式输出；WebSocket 全双工，适合实时交互，但连接治理更复杂。", "EduBot Chat/Live 用 SSE 映射不同事件帧。"),
            ("SSE 断线续传怎么做？", "事件带单调 ID，客户端回 Last-Event-ID/LGI，服务端从可重放缓冲继续；要防重复、跳帧和旧请求继续发流。", "Task LGI 与 Session 高水位要分开，incoming ACK 只能单调推进。"),
            ("gRPC 适合什么？", "内部强类型 RPC、HTTP/2 多路复用和双向流；不适合浏览器直连的普适性场景。需要处理 deadline、status、负载均衡和 protobuf 兼容。", "Agent Adapter 内部服务可考虑 gRPC，外部客户端仍可 SSE。"),
            ("API 网关做什么？", "认证鉴权、路由、限流、灰度、协议转换、熔断、观测和 WAF；不应承载复杂业务状态机。", "Router 是业务控制面，不等于 API Gateway。"),
            ("长连接的资源风险？", "每连接占 fd、内存、缓冲区和心跳；慢客户端会积压发送队列。需要连接上限、idle timeout、写超时、心跳和背压。", "Live/SSE 需要处理客户端断开并终止内部生成器。"),
        ]),
        ("六、Redis、缓存与多机房", [
            ("Cache Aside 怎么工作？", "读先查缓存，miss 回源并回填；写先更新数据库再删缓存是常见做法，但并发下仍需版本/延迟双删或消息驱动。", "主动服务的 pending slot 更像业务状态，不是普通 Cache Aside。"),
            ("Redis GETDEL/ Lua/CAS 的差别？", "GETDEL 原子取删适合一次消费但不支持 judge 后释放；Lua 可原子校验版本并转状态；WATCH/MULTI 是乐观事务，冲突需重试。", "steering 需要 CLAIMED lease，Lua 状态机比直接 GETDEL 更合适。"),
            ("缓存穿透、击穿、雪崩？", "穿透是查不存在 key，可布隆/空值；击穿是热点过期并发回源，可 singleflight；雪崩是大量同时过期，可随机 TTL、分批预热和降级。", "用户画像/会话缓存要分别处理不存在与热点。"),
            ("多机房为什么难做到强一致？", "跨区 RTT 高且网络分区不可避免；同步强一致影响可用性，异步复制有短暂重复/陈旧。要按业务选择主区、revision、幂等和冲突策略。", "主动服务后台多写、本地读是时延优先，必须接受短暂副本窗口。"),
            ("TTL 只是防内存泄漏吗？", "还表达业务新鲜度和租约边界。TTL 到期不等于逻辑已处理，关键状态还需终态/审计记录。", "pending steering TTL 防陈旧候选，但 Gate 仍需用最新 Query 重判。"),
        ]),
        ("七、Transformer、BERT 与分类训练", [
            ("Self-Attention 公式与复杂度？", "Attention(Q,K,V)=softmax(QK^T/sqrt(d_k))V。序列长度 n 下时间/显存通常 O(n^2)，多头让模型在不同子空间学习关系。", "解释长上下文为何昂贵，也支撑 token 预算设计。"),
            ("BERT 与 GPT 的区别？", "BERT 双向编码、MLM 预训练，适合理解/分类；GPT 自回归 next-token，适合生成。部署意图分类时 BERT 小模型可低延迟。", "高途 BERT 做 intent，LLM 处理低置信和回答生成。"),
            ("Focal Loss 是什么？", "FL=-(1-p_t)^gamma log p_t，降低高置信 easy sample 权重，突出 hard case；gamma=0 退化为 CE。", "与 class weight 解决不同维度，需消融。"),
            ("Label Smoothing 有什么用？", "把 one-hot 目标变为软分布，降低过度自信、提升泛化/校准；过大会伤害可分性。", "有助于 confidence threshold 更可信，但不能代替温度校准。"),
            ("类别不平衡怎么处理？", "重采样、class weight、Focal、数据增强、阈值按类设置、两阶段分类；评估用 macro/per-class F1 和 PR，而非只看 accuracy。", "课程类 33k vs 其他类 26 的极端不平衡要先治理数据。"),
            ("怎么做置信度校准？", "验证集画 reliability diagram、计算 ECE/Brier，做 temperature scaling；阈值按误判成本和 LLM 回退预算选。", "文本 0.85、图片 0.90 是配置，应通过校准和在线覆盖率验证。"),
            ("实体增强训练为何有效？", "把具体名字替换/标注为实体类型，模型学习结构信号而非记忆实体；训练推理必须用相同标注逻辑。", "高途新增老师只更新 entities.csv。"),
        ]),
        ("八、检索、RAG 与重排", [
            ("BM25 的核心思想？", "基于词频、逆文档频率并对文档长度归一化；词频收益会饱和，适合关键词精确匹配。", "与向量检索互补。"),
            ("稠密向量检索的优缺点？", "能召回语义相近表达，但对精确实体、数字和新词可能不稳定；还受 embedding 版本和 ANN 召回误差影响。", "高途用业务实体/intent Filter约束。"),
            ("混合检索怎么融合？", "可先分别召回，再用加权分数或 RRF 融合；原始分数量纲不同要归一化。之后用 cross-encoder/LLM 或业务特征重排。", "文档示例是向量+BM25，再做实体乘法 boost。"),
            ("RAG 的完整链路？", "切分、embedding、索引、Query 改写、召回、过滤/融合、重排、上下文拼装、生成、引用与评估。", "EventSummary 多、应按 Query RAG；FactMemory 少、可直接注入。"),
            ("chunk 怎么切？", "按语义/结构优先，大小与 overlap 在召回完整性和噪声/成本之间权衡；表格/对话要保结构和元数据。", "EventSummary 天然按 session，是很好的语义块。"),
            ("如何评估 RAG？", "分别评 retrieval Recall@K/MRR/NDCG、context precision/recall、answer faithfulness/accuracy 和端到端任务指标；建立无法回答集与引用核验。", "不要只看主观回答效果。"),
            ("为什么乘法 boost 可能有风险？", "多个实体连续相乘可能放大过度且分数失去可比性，需要上限、log-space 或学习排序；Filter OR 也可能扩大噪声。", "面试中主动给出离线 NDCG 与在线搜索成功率调参方案。"),
        ]),
        ("九、Context Engineering 与长期记忆", [
            ("上下文应该如何分层？", "system/策略、当前 Query、近期原文轮次、Topic/Task 状态、长期事实、按需召回历史各自有不同优先级和生命周期。", "高途近期保真+FactMemory+EventSummary RAG；阿里 Task-scoped context。"),
            ("为什么不能全量历史？", "成本和时延线性增长，旧信息干扰、隐私暴露、超窗截断不可控。应做预算、最近优先、摘要和检索。", "tiktoken 计数优先于字符估算。"),
            ("摘要会丢信息怎么办？", "原文保留权威存储，摘要带来源和时间；需要细节时 RAG 回源，关键 Task 状态不只存在自然语言摘要里。", "summary 不能替 checkpoint。"),
            ("长期记忆如何更新？", "先提取候选事实，再按类型/标签与旧事实做 ADD/SKIP/MERGE/REPLACE；保留 inactive 历史和来源，处理 TTL、删除与用户纠正。", "高途 FactMemory 协议。"),
            ("什么信息不该记？", "临时闲聊、敏感信息、未经确认推断、短期情绪和无长期价值内容；应按用途、同意、保留期与可删除性治理。", "教育未成年人场景尤其严格。"),
            ("Prompt injection 如何防？", "外部内容与系统指令分层，工具/控制信号结构化并白名单，检索内容视为数据，限制工具权限与输出校验。", "用户文本不能伪造 Handoff/Handback/steering。"),
        ]),
        ("十、多模态、视频与 ASR", [
            ("视频为什么常转成帧+音频文本？", "便于复用图文 MLLM 接口、控制输入成本，并把视觉和语言证据显式化；代价是时序信息损失。", "高途 20 帧+ASR。"),
            ("均匀抽帧与关键帧方法对比？", "均匀简单稳定但会漏短镜头；镜头切分覆盖场景变化；CLIP/熵/字幕变化可选高信息帧，但计算更贵。", "可先均匀基线再做离线消融。"),
            ("ASR 常见问题？", "口音、噪声、专有词、断句、时间戳和幻觉；可用 VAD、热词、语言模型、置信度和字幕/视觉交叉验证。", "音频文本在教育视频中常是核心证据。"),
            ("多标签分类怎么评？", "micro/macro F1、per-label precision/recall、Precision@K、subset accuracy、空标签正确率和标签一致性；类别长尾要分层。", "兴趣和地域应分别评估。"),
            ("结构化输出为什么仍需解析/校验？", "LLM 可能返回非 JSON、自创字段/标签或越界值。要 schema 校验、候选白名单、重试/repair 和业务后验规则。", "高途只保留 tag 映射中存在的结果。"),
        ]),
        ("十一、可观测性与实验", [
            ("Metrics、Logs、Traces 怎么分工？", "Metrics 看趋势和告警，Logs 解释单事件细节，Trace 串跨服务关键路径。高基数 ID 放日志/Trace，不做 metric label。", "EduBot 以 request/session/task/intervention 关联。"),
            ("为什么 HTTP 200 不等于业务成功？", "它只证明协议请求成功；还要验证配置命中、算法判断、状态 apply、下游执行、最终 responder、客户端接收和持久化。", "两个阿里项目都强调分阶段语义验收。"),
            ("A/B 的 p-value 表示什么？", "在零假设下观察到当前或更极端差异的概率，不是‘方案为真的概率’，也不表示效果大小。要同时报告 uplift、置信区间与业务意义。", "主动服务 p≤0.0096 说明差异显著，但仍需归因口径。"),
            ("意向治疗与实际曝光分析区别？", "ITT 按分桶比较，保留随机化但会被低曝光稀释；treated-only 按实际曝光可能有选择偏差。可用触发资格、工具变量或受控曝光设计。", "主动服务只有 2.57% 实际注入，不能直接把 ITT lift 当单次注入效应。"),
            ("多指标实验如何防误判？", "预注册主指标和 guardrail，控制多重比较，检查样本比例失衡、随机化、时序、机器人/拨测流量和新奇效应。", "路由单日 1128/197 样本应谨慎。"),
            ("如何设计审计抽样？", "链路健康度尽量全量；语义质量按关键分支定向覆盖，再补随机样本估总体。异常逐个下钻，所有结论保留分母和边界。", "阿里审计先 902 请求全量，再 18 Session 定向语义。"),
        ]),
        ("十二、安全、隐私与降级", [
            ("fail-open 和 fail-closed 怎么选？", "核心业务可用性和附加能力风险分别判断。主回答可以 fail-open 继续，主动干预应 fail-closed 不展示未经确认内容。", "Handoff Gate 已知 active 时继续；Proactive 异常返回空。"),
            ("最小权限如何落地？", "服务/工具只获取完成任务所需数据和动作，凭据独立、短期、可轮换；用户/Session/Task 关联要鉴权，日志脱敏。", "Agent 工具白名单和 Task-scoped context。"),
            ("如何治理未成年人数据？", "用途限制、最小化、敏感字段隔离、短保留期、访问审计、可删除/纠正，避免用情绪或健康信息做不必要推荐。", "主动服务的画像/情绪处理风险更高。"),
            ("降级为什么要分类？", "disabled、skip、miss、timeout、error、stale、conflict、delivered 的业务含义不同；混在一起会掩盖依赖故障或误判。", "每次 fallback 都要有 reason。"),
            ("如何防内部控制协议被用户伪造？", "控制消息使用结构化类型、签名/来源标识、白名单字段和服务端生成 ID；用户文本永远按数据处理，不能直接改变 owner 或 Task 状态。", "Handoff/Handback/steering 不接受自然语言伪造。"),
        ]),
    ]

    for chapter, items in chapters:
        add_heading(doc, chapter, 1)
        qa = []
        for q, a, mapping in items:
            follows = ["项目映射：" + mapping] if mapping else None
            qa.append((q, a, follows))
        add_qa(doc, qa)

    add_heading(doc, "十三、系统设计题答题模板", 1)
    add_steps(doc, [
        "先问清规模与 SLA：QPS、并发连接、消息长度、首 token/p99、可用性、数据新鲜度、是否允许重复。",
        "定义业务身份和不变量：request/session/topic/task/intervention；谁是 owner；什么必须单调或只能一次。",
        "画主链和异步链：入口/网关 -> 控制面 -> 执行面 -> 存储/缓存 -> 流式交付 -> 观测。",
        "分读写路径与状态：权威存储、缓存、版本、TTL、幂等键、outbox、重试和补偿。",
        "做容量估算：峰值 QPS、连接数、消息大小、缓存/日志量、模型吞吐和并发预算。",
        "设计失败语义：每个依赖超时后回什么、状态是否提交、是否释放 claim、是否允许重试。",
        "说明观测与实验：阶段事件、p50/p95/p99、质量指标、业务指标、guardrail、灰度和回滚。",
        "最后主动讲边界：一期简化、未完成能力、P0 风险和演进顺序。",
    ])

    add_heading(doc, "十四、14 天复习节奏", 1)
    add_table(doc, ["天", "主题", "验收"], [
        ["1-2", "阿里 Router/Task 项目 + 状态机", "能白板画五动作，2 分钟讲清 advice/effect"],
        ["3", "CAS/幂等/outbox/lease", "用四个故障窗口解释各机制"],
        ["4", "SSE/连接池/超时重试/背压", "完成一个 LLM 流式服务系统设计"],
        ["5-6", "高途搜索 + BERT", "手写 Attention/Focal，讲清实体增强和阈值"],
        ["7", "BM25/向量/RAG/重排", "能设计离线+在线评估"],
        ["8", "用户画像与长期记忆", "讲清三层记忆、merge、TTL、隐私"],
        ["9", "多模态标签/ASR/抽帧", "讲清消融、评价指标和工程降级"],
        ["10", "主动服务", "画 prepare-confirm、L0-L4、atomic claim"],
        ["11", "A/B 与可观测性", "正确解释 p-value、ITT、2.57% 注入"],
        ["12", "简历逐条反问", "每个数字能说来源、口径、边界"],
        ["13", "模拟一面", "项目 2 分钟 + 每个项目 10 个追问"],
        ["14", "模拟二面/系统设计", "45 分钟完成设计并复盘取舍"],
    ], [900, 4000, GUIDE_TABLE_WIDTH - 4900])

    add_heading(doc, "十五、最后的表达红线", 1)
    add_bullets(doc, [
        "不把目标设计说成已上线；不把 Demo 说成完整生产系统。",
        "不把估算时延说成个人压测；不为高途项目补不存在的准确率/收益。",
        "不隐去样本量、时间窗口、实际干预率和抽样方式。",
        "不把 HTTP 200、Gate hit、缓存命中当用户效果。",
        "回答技术债时先承认事实，再说影响、定位证据和改进顺序。",
        "所有项目至少准备：30 秒版本、2 分钟版本、架构图、一个难点、一个故障、一个权衡、一个未完成边界。",
    ])

    add_sources(doc, [
        "高途实习材料/阚海_简历_27应届.pdf",
        "高途实习材料/TECH_aisearch.md",
        "高途实习材料/TECH_userprofile.md",
        "高途实习材料/tagdoc.md",
        "阿里实习材料/01-路由与Handoff完整业务逻辑.md",
        "阿里实习材料/02-任务编排与Topic完整设计.md",
        "阿里实习材料/03-主动服务完整业务与算法设计.md",
        "阿里实习材料/04-统一架构与链路工程设计.md",
        "阿里实习材料/edubot-on-bucket-audit-2026-08-26.md",
        "阿里实习材料/各类实验收益.txt 与主动服务实验截图",
    ], "本手册按你的项目反向组织八股；定义性内容可作为知识储备，涉及项目状态和数字时仍以对应项目文档的证据边界为准。")
    return save_doc(doc, "06_技术栈八股与项目追问手册.docx")


def main() -> None:
    paths = [
        build_resume(),
        build_ali_router(),
        build_ali_proactive(),
        build_gaotu_search(),
        build_gaotu_profile(),
        build_gaotu_tags(),
        build_tech_handbook(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
